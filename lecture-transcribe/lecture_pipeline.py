from __future__ import annotations

"""
lecture_pipeline.py v3 — 錄音 → 摘要 → Notion 完整 Pipeline

功能:
  - 雙 DB 路由 (EMBA 課堂 / 商務會談)
  - 四種摘要模板自動切換
  - Speaker 兩階段識別 (預填 + 校正)
  - 多段錄音合併
  - 長音訊自動切割
  - 逐字稿 toggle heading 收合
  - LLM 模型可指定

入口: handle_audio_message() — 由 OpenClaw Telegram handler 呼叫
"""

import html
import json
import os
import time
import urllib.parse
import urllib.request
import logging
import re
import subprocess
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple

import yaml

from transcribe import (
    transcribe_audio, run_diarization, get_audio_duration,
    merge_audio_files, get_speaker_preview,
    build_speaker_quotes, filter_non_topic_speakers,
)
from summary_prompts import (
    select_template, select_model, select_chunk_model,
    build_summary_prompt, build_system_prompt, build_chunk_prompt, build_reduce_prompt,
    TEMPLATE_NAMES, apply_course_specific_summary_flags,
)
from notion_upload import upload_emba, upload_business

logger = logging.getLogger(__name__)

CONFIG = {
    "schedule_path": Path(__file__).parent / "course_schedule.yaml",
    "output_dir": Path.home() / "whisperx-outputs",
    # Cache transcripts so a gateway restart or retry doesn't re-transcribe long audio.
    # (Safe: stored locally under the output dir only.)
    "cache_dir": Path.home() / "whisperx-outputs" / "cache",
    "reference_dir": Path.home() / "whisperx-outputs" / "references",
    "max_reference_chars_per_source": 6000,
    "max_reference_chars_total": 18000,
    "max_ocr_reference_chars_per_source": 9000,
}

def _env_flag(name: str, default: bool = False) -> bool:
    raw = (os.environ.get(name) or '').strip().lower()
    if not raw:
        return default
    return raw in {'1', 'true', 'yes', 'on'}


ALLOW_GOOGLE_CHUNK_FALLBACK = _env_flag('LECTURE_TRANSCRIBE_ALLOW_GOOGLE_CHUNK_FALLBACK', False)
ALLOW_GOOGLE_FINAL_FALLBACK = _env_flag('LECTURE_TRANSCRIBE_ALLOW_GOOGLE_FINAL_FALLBACK', False)
ALLOW_GOOGLE_DIRECT_FALLBACK = _env_flag('LECTURE_TRANSCRIBE_ALLOW_GOOGLE_DIRECT_FALLBACK', False)


def _build_chunk_fallback_chain() -> list[str]:
    chain = [
        "google-ai/gemma-4-31b-it",
        "openai-codex/gpt-5.4",
    ]
    if ALLOW_GOOGLE_CHUNK_FALLBACK:
        chain.append("google/gemini-2.5-flash")
    return chain


def _build_final_fallback_chain() -> list[str]:
    chain = [
        "minimax-portal/MiniMax-M2.7",
    ]
    if ALLOW_GOOGLE_FINAL_FALLBACK:
        chain.append("google/gemini-3-pro-preview")
    return chain


LLM_FINAL_FALLBACK_CHAIN = _build_final_fallback_chain()
LLM_CHUNK_FALLBACK_CHAIN = _build_chunk_fallback_chain()

_GEMINI_DIRECT_MODEL_MAP = {
    'google/gemini-2.5-flash': 'gemini-2.5-flash',
    'google/gemini-2.5-pro': 'gemini-2.5-pro',
    'google/gemini-3-pro-preview': 'gemini-2.5-pro',
    'google/gemini-3.1-pro-preview': 'gemini-2.5-pro',
    'google/gemini-3.1-flash-lite-preview': 'gemini-2.5-flash-lite',
    'google/gemini-3-flash-preview': 'gemini-2.5-flash',
}


def _append_proc_warning(proc: Optional[dict], message: str) -> None:
    if not isinstance(proc, dict) or not message:
        return
    warnings = proc.setdefault('llm_warnings', [])
    if message not in warnings:
        warnings.append(message)


def _describe_llm_policy() -> dict:
    return {
        'chunk_default': 'minimax-portal/MiniMax-M2.7',
        'chunk_fallback_chain': LLM_CHUNK_FALLBACK_CHAIN[:],
        'final_fallback_chain': LLM_FINAL_FALLBACK_CHAIN[:],
        'allow_google_chunk_fallback': ALLOW_GOOGLE_CHUNK_FALLBACK,
        'allow_google_final_fallback': ALLOW_GOOGLE_FINAL_FALLBACK,
        'allow_google_direct_fallback': ALLOW_GOOGLE_DIRECT_FALLBACK,
    }



# ============================================================
# Markdown export / Obsidian archive
# ============================================================

def _safe_note_name(metadata: dict) -> str:
    name = (metadata.get('course_name') or metadata.get('meeting_name') or 'rec')
    return name.replace('/', '-').replace(' ', '_')


def _resolve_obsidian_vault_path() -> Optional[Path]:
    env_path = (os.environ.get('OBSIDIAN_VAULT_PATH') or '').strip()
    if env_path:
        p = Path(env_path).expanduser()
        if p.exists():
            return p

    try:
        res = subprocess.run(
            ['obsidian-cli', 'print-default', '--path-only'],
            capture_output=True,
            text=True,
            timeout=10,
            check=True,
        )
        p = Path((res.stdout or '').strip()).expanduser()
        if p.exists():
            return p
    except Exception:
        pass

    cfg = Path.home() / 'Library' / 'Application Support' / 'obsidian' / 'obsidian.json'
    try:
        import json
        data = json.loads(cfg.read_text(encoding='utf-8'))
        for v in (data.get('vaults') or {}).values():
            if v.get('open') and v.get('path'):
                p = Path(v['path']).expanduser()
                if p.exists():
                    return p
    except Exception:
        pass
    return None


def _obsidian_note_path(metadata: dict, vault_path: Optional[Path]) -> Optional[Path]:
    if not vault_path:
        return None

    date = metadata.get('date', 'x')
    name = _safe_note_name(metadata)
    top = 'EMBA' if metadata.get('type') == 'emba' else '商務會談'
    course = metadata.get('course_name') or metadata.get('meeting_name') or '未分類'
    if metadata.get('type') == 'emba':
        folder = vault_path / top / '02 每週課堂筆記' / course
    else:
        folder = vault_path / top / course / '課堂筆記'
    return folder / f"{date}_{name}.md"


def _strip_emba_transcript_sections(summary_md: str, metadata: dict) -> str:
    """EMBA 模板不應輸出完整逐字稿；若模型誤帶出相關章節就裁掉。"""
    if metadata.get('type') != 'emba':
        return summary_md

    text = summary_md or ''
    patterns = [
        r'\n###\s*完整逐字稿[\s\S]*$',
        r'\n##\s*完整逐字稿[\s\S]*$',
        r'\n###\s*逐字稿[\s\S]*$',
        r'\n##\s*逐字稿[\s\S]*$',
    ]
    for pat in patterns:
        text = re.sub(pat, '', text, flags=re.I)
    return text.rstrip() + '\n'


def _build_notion_footer_blocks(metadata: dict, obsidian_md_path: Optional[Path] = None) -> list:
    """Build Notion API block children for the unified footer section.

    Uses custom_emoji mentions (from icons.py) for brand icons in headings and model lines.
    All bullet text is italic + gray color. Applies to ALL templates.
    """
    from icons import notion_emoji_mention, notion_model_emoji

    proc = metadata.get('process_summary') or {}
    llm_usage = proc.get('llm_usage') or {}
    chunk_usage = llm_usage.get('chunk') or {}
    final_usage = llm_usage.get('final') or {}

    def _it(content: str, color: str = 'gray') -> dict:
        """Italic text block with specified color."""
        return {"type": "text", "text": {"content": content},
                "annotations": {"bold": False, "italic": True, "strikethrough": False,
                                "underline": False, "code": False, "color": color}}

    def _plain(content: str) -> dict:
        return {"type": "text", "text": {"content": content},
                "annotations": {"bold": False, "italic": False, "strikethrough": False,
                                "underline": False, "code": False, "color": "default"}}

    def bullet(*parts, children=None, color='default'):
        b = {"object": "block", "type": "bulleted_list_item",
             "bulleted_list_item": {"rich_text": [p for p in parts if p], "color": color}}
        if children:
            b["bulleted_list_item"]["children"] = children
        return b

    def heading3(*parts):
        return {"object": "block", "type": "heading_3",
                "heading_3": {"rich_text": [p for p in parts if p]}}

    def toggle_heading3(*parts, children=None):
        block = {"object": "block", "type": "heading_3",
                 "heading_3": {"rich_text": [p for p in parts if p], "is_toggleable": True}}
        if children:
            block["heading_3"]["children"] = children
        return block

    def divider():
        return {"object": "block", "type": "divider", "divider": {}}

    def quote(*parts):
        return {"object": "block", "type": "quote",
                "quote": {"rich_text": [p for p in parts if p]}}

    def paragraph():
        return {"object": "block", "type": "paragraph", "paragraph": {"rich_text": []}}

    blocks = []
    footer_children = []

    # ── Heading with OpenClaw emoji ──
    oc_emoji = notion_emoji_mention('openclaw-dark')

    # ── 來源取得 ──
    source = '雲端分享音檔下載後處理' if proc.get('downloaded_from_cloud') else '本機/既有音檔直接處理'
    footer_children.append(bullet(_it(f'來源取得：{source}')))

    # ── 參考資料整合 ──
    ref_sources = proc.get('reference_sources') or []
    if ref_sources:
        src_list = '、'.join(ref_sources[:6])
        more = f' 等 {len(ref_sources)} 項' if len(ref_sources) > 6 else f'（共 {len(ref_sources)} 項）'
        footer_children.append(bullet(_it(f'參考資料整合：本次納入參考來源 {more}：{src_list}，並與逐字稿交叉校對後再整理摘要。')))
    else:
        footer_children.append(bullet(_it('參考資料整合：本次為單次錄音處理，未納入額外參考資料。')))

    if proc.get('skip_notion_ai_transcript'):
        if proc.get('notion_ai_transcript_fallback_used'):
            footer_children.append(bullet(_it('Notion AI 逐字稿策略：已啟用 skip_notion_ai_transcript；原始錄音轉錄品質判定偏差，因此改以 Notion AI 逐字稿做輔助校對 fallback。')))
        else:
            footer_children.append(bullet(_it('Notion AI 逐字稿策略：已啟用 skip_notion_ai_transcript；Notion AI 長逐字稿預設只當 reference，不作主逐字稿來源。')))

    # ── 音訊前處理 ──
    preprocess_tool = proc.get('preprocess_tool') or 'ffmpeg'
    preprocess_desc = proc.get('preprocess_desc') or '16kHz 單聲道正規化與 loudnorm 音量校正'
    footer_children.append(bullet(_it(f'音訊前處理：使用 {preprocess_tool} 進行 {preprocess_desc}，讓後續辨識更穩定。')))

    # ── 逐字稿轉錄 ──
    transcribe_engine = proc.get('transcribe_engine') or 'whisper'
    seg_count = proc.get('segment_count')
    dur_text = proc.get('audio_duration_text') or ''
    t_line = f'逐字稿轉錄：使用 {transcribe_engine} 產出逐字稿'
    if dur_text:
        t_line += f'（音訊長度約 {dur_text}）'
    if seg_count:
        t_line += f'，共切出約 {seg_count} 個語意片段'
    t_line += '。'
    footer_children.append(bullet(_it(t_line)))

    # ── 快取策略 ──
    if proc.get('used_cache'):
        footer_children.append(bullet(_it('快取策略：本次有命中既有轉錄快取，略過重跑逐字稿。')))
    else:
        footer_children.append(bullet(_it('快取策略：本次未命中既有轉錄快取，逐字稿為重新轉錄產生。')))

    # ── 摘要生成 + 模型資訊 ──
    sub_items = []

    if proc.get('used_chunking'):
        chunk_count = proc.get('chunk_count') or '多'
        chunk_model_req = proc.get('chunk_model') or chunk_usage.get('requested_model') or '輕量模型'
        final_model_req = proc.get('final_model') or final_usage.get('requested_model') or '主摘要模型'
        summary_desc = f'摘要生成：長音檔模式，先分成 {chunk_count} 段做分段重點整理，再由主摘要模型整併。使用模型如下。'

        # Chunk model items
        chunk_emoji = notion_model_emoji(chunk_model_req)
        sub_items.append(bullet(chunk_emoji, _it(f' Chunk 指定模型：{chunk_model_req}')))
        chunk_success_list = chunk_usage.get('success_models') or []
        if chunk_success_list:
            fb_tag = '（有觸發 fallback）' if chunk_usage.get('used_fallback') else '（未觸發 fallback）'
            success_str = '、'.join(chunk_success_list)
            first_emoji = notion_model_emoji(chunk_success_list[0])
            sub_items.append(bullet(first_emoji, _it(f' Chunk 實際成功模型：{success_str} {fb_tag}')))
        chunk_chain_list = proc.get('chunk_fallback_chain') or []
        if chunk_chain_list:
            chain_str = ' → '.join(chunk_chain_list)
            sub_items.append(bullet(_it(f'Chunk fallback 鏈：{chain_str}')))
    else:
        final_model_req = proc.get('final_model') or final_usage.get('requested_model') or '主摘要模型'
        summary_desc = '摘要生成：直接對完整逐字稿生成摘要，使用模型如下。'

    # Final model items
    final_emoji = notion_model_emoji(final_model_req)
    sub_items.append(bullet(final_emoji, _it(f' Final 指定模型：{final_model_req}')))
    final_success = final_usage.get('final_used_model') or ''
    if final_success:
        fb_tag = '（有觸發 fallback）' if final_usage.get('used_fallback') else '（未觸發 fallback）'
        success_emoji = notion_model_emoji(final_success)
        sub_items.append(bullet(success_emoji, _it(f' Final 實際成功模型：{final_success} {fb_tag}')))
    final_chain_list = final_usage.get('fallback_chain') or proc.get('final_fallback_chain') or []
    if final_chain_list:
        # Build emoji-prefixed chain items: :gemini-color: google/gemini-3-pro-preview → :minimax-color: minimax-portal/MiniMax-M2.7 → ...
        chain_parts = []
        for model_id in final_chain_list:
            emoji = notion_model_emoji(model_id)
            if emoji:
                chain_parts.append(emoji)
                chain_parts.append(_it(f' {model_id}'))
            else:
                chain_parts.append(_it(f' {model_id}'))
            if model_id != final_chain_list[-1]:
                chain_parts.append(_it(' → '))
        sub_items.append(bullet(_it('Final fallback 鏈：'), *chain_parts))

    policy = proc.get('llm_policy') or {}
    if not policy.get('allow_google_direct_fallback'):
        sub_items.append(bullet(_it('Guardrail：Gemini direct fallback 預設停用；只有明確開啟環境旗標才會直連 Google API。')))

    for warning in proc.get('llm_warnings') or []:
        sub_items.append(bullet(_it(warning, color='orange')))

    footer_children.append(bullet(_it(summary_desc), children=sub_items))

    # ── 說話者處理 ──
    diarization = '有啟用說話者分辨（speaker diarization）' if proc.get('used_diarization') else '未啟用說話者分辨 / speaker diarization（以穩定與速度優先）'
    footer_children.append(bullet(_it(f'說話者處理：{diarization}。')))

    ocr_summary = proc.get('ocr_summary') or {}
    if ocr_summary.get('image_count'):
        ocr_line = (
            f"OCR 補充：共處理 {ocr_summary.get('image_count', 0)} 張圖片，"
            f"成功 {ocr_summary.get('success_count', 0)} 張，失敗 {ocr_summary.get('failed_count', 0)} 張，"
            f"擷取約 {ocr_summary.get('total_chars', 0)} 字。"
        )
        if ocr_summary.get('warning'):
            ocr_line += f" 注意：{ocr_summary.get('warning')}。"
        footer_children.append(bullet(_it(ocr_line)))
        if metadata.get('detailed_material_restore'):
            footer_children.append(bullet(_it('教材還原模式：已啟用 detailed_material_restore / high_priority_exam_material，教材照片 OCR 會以高保留度方式納入最終摘要。')))

    # ── 耗時 ──
    transcribe_elapsed = proc.get('transcribe_elapsed_text') or ''
    chunk_elapsed = proc.get('chunk_elapsed_text') or ''
    summary_elapsed = proc.get('summary_elapsed_text') or ''
    total_elapsed = proc.get('total_elapsed_text') or ''
    transcribe_sec = proc.get('transcribe_elapsed_sec') or 0
    timing_parts = []
    if transcribe_elapsed:
        timing_parts.append(f'轉錄 {transcribe_elapsed}')
    if chunk_elapsed:
        timing_parts.append(f'分段整理 {chunk_elapsed}')
    if summary_elapsed:
        timing_parts.append(f'摘要生成 {summary_elapsed}')
    if timing_parts and total_elapsed:
        timing_str = f"耗時：{' ＋ '.join(timing_parts)}，總計約 {total_elapsed}。"
    elif timing_parts:
        timing_str = f"耗時：{' ＋ '.join(timing_parts)}。"
    elif total_elapsed:
        timing_str = f'耗時：總計約 {total_elapsed}。'
    elif transcribe_sec > 0:
        timing_str = f'耗時：轉錄約 {int(transcribe_sec)} 秒。'
    else:
        timing_str = '耗時：未記錄詳細分段耗時。'
    footer_children.append(bullet(_it(timing_str)))

    # ── 上傳歸檔 ──
    notion_emoji = notion_emoji_mention('notion')
    oc_emoji = notion_emoji_mention('openclaw-dark')
    notion_label = 'Notion 課堂摘要庫' if metadata.get('type') == 'emba' else 'Notion 商務會談摘要 DB'
    footer_children.append(bullet(
        _it('上傳歸檔：同步寫入 '),
        notion_emoji, _it(f' {notion_label}'),
        _it(' ＋ '), oc_emoji, _it(' Obsidian 本地 Vault。'),
    ))

    # ── 免責聲明 ──
    footer_children.append(paragraph())
    footer_children.append(quote(_it(
        '⚠️ 本摘要由大語言模型（LLM）依據錄音轉錄內容、使用者筆記及其他補充材料交叉校對後整理產出，'
        '仍可能存在誤解、遺漏或表述偏差；若需正式引用、對外使用或作為決策依據，請務必加以查核。',
        color='gray',
    )))

    blocks.append(toggle_heading3(oc_emoji, _plain('作業路徑說明'), children=footer_children))

    # ── Obsidian 知識庫連結 ──
    blocks.append(divider())
    blocks.append(heading3(_plain(':obsidian-color: Obsidian 知識庫連結')))
    if obsidian_md_path:
        # Build obsidian:// deep link: vault=OpenClaw&file=EMBA/02 每週課堂筆記/跨文化交流與研習/...
        vault_rel = str(obsidian_md_path).split('小龍女知識庫/')[-1] if '小龍女知識庫/' in str(obsidian_md_path) else str(obsidian_md_path.name)
        # Remove .md extension for obsidian URL
        vault_file = vault_rel.replace('.md', '')
        oc_url = f'obsidian://open?vault=OpenClaw&file={urllib.parse.quote(vault_file)}'
        blocks.append(bullet(
            _plain('📓 本堂課堂筆記　'),
            {"type": "text", "text": {"content": vault_rel, "link": {"url": oc_url}},
             "annotations": {"bold": False, "italic": False, "strikethrough": False, "underline": False, "code": False, "color": "default"}}
        ))
    course = metadata.get('course_name') or metadata.get('meeting_name') or ''
    if course:
        blocks.append(bullet(_plain(f'📚 課程主檔　[[{course}]]')))
    blocks.append(bullet(_plain('📅 學期總覽　[[114-2 課程總覽]]')))

    return blocks


def _build_model_runtime_lines(proc: dict) -> list[str]:
    """Build compact model info lines for the 'done' notification message."""
    from icons import model_tag
    proc = proc or {}
    usage = proc.get('llm_usage') or {}
    chunk_usage = usage.get('chunk') or {}
    final_usage = usage.get('final') or {}
    lines = []

    chunk_models = chunk_usage.get('success_models') or []
    if chunk_models:
        tags = [f'{model_tag(m)} {m}' if model_tag(m) else m for m in chunk_models]
        prefix = 'Chunk 實際模型'
        if chunk_usage.get('used_fallback'):
            prefix += '（含 fallback）'
        lines.append(f"⚙️ {prefix}：{'、'.join(tags)}")

    final_model = final_usage.get('final_used_model') or ''
    if final_model:
        tag = model_tag(final_model)
        prefix = 'Final 實際模型'
        if final_usage.get('used_fallback'):
            prefix += '（含 fallback）'
        model_str = f'{tag} {final_model}' if tag else final_model
        lines.append(f"🧠 {prefix}：{model_str}")

    for warning in proc.get('llm_warnings') or []:
        lines.append(warning)

    return lines


def _append_summary_footer(summary_md: str, metadata: dict, audio_path: str, cache_path: Path, md_path: Path, obsidian_md_path: Optional[Path]) -> str:
    """Append unified pipeline footer to summary markdown.

    Format applies to ALL templates (A/B/C/D1/D2/D3/E/W/R).
    Uses [Brand] text tags for model identification in Markdown output.
    Notion output uses custom_emoji mentions (built separately by _build_notion_footer_blocks).
    """
    from icons import model_emoji_tag, service_emoji_tag

    proc = metadata.get('process_summary') or {}
    parts = [summary_md.rstrip(), '', '---', '', '### :openclaw-dark:作業路徑說明']

    # ── 來源取得 ──
    cloud_src_map = {'gdrive': 'Google Drive', 'onedrive': 'OneDrive'}
    src_label = cloud_src_map.get(proc.get('cloud_source', '').lower(), '雲端')
    audio_fmt = proc.get('audio_format') or metadata.get('audio_format') or ''
    fmt_str = f'（{audio_fmt} 格式）' if audio_fmt else ''
    if proc.get('downloaded_from_cloud'):
        source = f'{src_label} 雲端 .{audio_fmt or "音訊"} 檔下載後處理'
    else:
        source = f'本機 / 既有音訊檔{fmt_str}直接處理'
    parts.append(f"- *來源取得：{source}*")

    # ── 參考資料整合（有用到外部材料才顯示）──
    ref_sources = proc.get('reference_sources') or []
    if ref_sources:
        src_list = '、'.join(ref_sources[:6])
        more = f' 等 {len(ref_sources)} 項' if len(ref_sources) > 6 else f'（共 {len(ref_sources)} 項）'
        parts.append(f"- *參考資料整合：本次納入參考來源 {more}：{src_list}，並與逐字稿交叉校對後再整理摘要。*")
    else:
        parts.append('- *參考資料整合：本次為單次錄音處理，未納入額外參考資料。*')

    if proc.get('skip_notion_ai_transcript'):
        if proc.get('notion_ai_transcript_fallback_used'):
            parts.append('- *Notion AI 逐字稿策略：已啟用 skip_notion_ai_transcript；原始錄音轉錄品質判定偏差，因此改以 Notion AI 逐字稿做輔助校對 fallback。*')
        else:
            parts.append('- *Notion AI 逐字稿策略：已啟用 skip_notion_ai_transcript；Notion AI 長逐字稿預設只當 reference，不作主逐字稿來源。*')

    # ── 音訊前處理 ──
    preprocess_tool = proc.get('preprocess_tool') or 'ffmpeg'
    preprocess_desc = proc.get('preprocess_desc') or '16kHz 單聲道正規化與 loudnorm 音量校正'
    parts.append(f"- *音訊前處理：使用 {preprocess_tool} 進行 {preprocess_desc}，讓後續辨識更穩定。*")

    # ── 逐字稿轉錄 ──
    transcribe_engine = proc.get('transcribe_engine') or 'whisper'
    seg_count = proc.get('segment_count')
    dur_text = proc.get('audio_duration_text') or ''
    transcribe_line = f"*逐字稿轉錄：使用 {transcribe_engine} 產出逐字稿"
    if dur_text:
        transcribe_line += f"（音訊長度約 {dur_text}）"
    if seg_count:
        transcribe_line += f"，共切出約 {seg_count} 個語意片段"
    transcribe_line += '。*'
    parts.append(f"- {transcribe_line}")

    # ── 快取策略 ──
    if proc.get('used_cache'):
        parts.append('- *快取策略：本次有命中既有轉錄快取，因此略過重跑逐字稿，整體處理時間較短。*')
    else:
        parts.append('- *快取策略：本次未命中既有轉錄快取，逐字稿為重新轉錄產生。*')

    # ── 摘要生成 + 模型資訊 ──
    llm_usage = proc.get('llm_usage') or {}
    chunk_usage = llm_usage.get('chunk') or {}
    final_usage = llm_usage.get('final') or {}

    if proc.get('used_chunking'):
        chunk_count = proc.get('chunk_count') or '多'
        chunk_model_req = proc.get('chunk_model') or chunk_usage.get('requested_model') or '輕量模型'
        final_model_req = proc.get('final_model') or final_usage.get('requested_model') or '主摘要模型'
        parts.append(f"- *摘要生成：長音檔模式，先分成 {chunk_count} 段做分段重點整理，再由主摘要模型整併。*")

        # Chunk 模型（子項）
        chunk_tag = model_emoji_tag(chunk_model_req)
        parts.append(f"  - *{chunk_tag} Chunk 指定模型：{chunk_model_req}*")
        chunk_success_list = chunk_usage.get('success_models') or []
        if chunk_success_list:
            chunk_success_tags = [f'{model_emoji_tag(m)} {m}' if model_emoji_tag(m) else m for m in chunk_success_list]
            chunk_success = '、'.join(chunk_success_tags)
            fb_tag = '（有觸發 fallback）' if chunk_usage.get('used_fallback') else '（未觸發 fallback）'
            parts.append(f"  - *Chunk 實際成功模型：{chunk_success} {fb_tag}*")
        chunk_chain = ' → '.join(proc.get('chunk_fallback_chain') or [])
        if chunk_chain:
            parts.append(f"  - *Chunk fallback 鏈：{chunk_chain}*")
    else:
        final_model_req = proc.get('final_model') or final_usage.get('requested_model') or '主摘要模型'
        parts.append(f"- *摘要生成：直接對完整逐字稿生成摘要，使用模型如下。*")

    # Final 模型（子項）
    final_tag = model_emoji_tag(final_model_req)
    parts.append(f"  - *{final_tag} Final 指定模型：{final_model_req}*")
    final_success = final_usage.get('final_used_model') or ''
    if final_success:
        fb_tag = '（有觸發 fallback）' if final_usage.get('used_fallback') else '（未觸發 fallback）'
        success_tag = model_emoji_tag(final_success)
        parts.append(f"  - *{success_tag} Final 實際成功模型：{final_success} {fb_tag}*")
    final_chain = ' → '.join(
        final_usage.get('fallback_chain') or proc.get('final_fallback_chain') or []
    )
    if final_chain:
        parts.append(f"  - *Final fallback 鏈：{final_chain}*")

    policy = proc.get('llm_policy') or {}
    if not policy.get('allow_google_direct_fallback'):
        parts.append('- *Guardrail：Gemini direct fallback 預設停用；只有明確開啟環境旗標才會直連 Google API。*')
    for warning in proc.get('llm_warnings') or []:
        parts.append(f"- *{warning}*")

    # ── 說話者處理 ──
    diarization = '有啟用說話者分辨（speaker diarization）' if proc.get('used_diarization') else '未啟用說話者分辨 / speaker diarization（以穩定與速度優先）'
    parts.append(f"- *說話者處理：{diarization}。*")

    ocr_summary = proc.get('ocr_summary') or {}
    if ocr_summary.get('image_count'):
        ocr_line = (
            f"OCR 補充：共處理 {ocr_summary.get('image_count', 0)} 張圖片，"
            f"成功 {ocr_summary.get('success_count', 0)} 張，失敗 {ocr_summary.get('failed_count', 0)} 張，"
            f"擷取約 {ocr_summary.get('total_chars', 0)} 字。"
        )
        if ocr_summary.get('warning'):
            ocr_line += f" 注意：{ocr_summary.get('warning')}。"
        parts.append(f"- *{ocr_line}*")
        if metadata.get('detailed_material_restore'):
            parts.append('- *教材還原模式：已啟用 detailed_material_restore / high_priority_exam_material，教材照片 OCR 會以高保留度方式納入最終摘要。*')

    # ── 耗時 ──
    transcribe_elapsed = proc.get('transcribe_elapsed_text') or ''
    transcribe_sec = proc.get('transcribe_elapsed_sec') or 0
    chunk_elapsed = proc.get('chunk_elapsed_text') or ''
    summary_elapsed = proc.get('summary_elapsed_text') or ''
    total_elapsed = proc.get('total_elapsed_text') or ''

    timing_parts = []
    if transcribe_elapsed:
        timing_parts.append(f"轉錄 {transcribe_elapsed}")
    if chunk_elapsed:
        timing_parts.append(f"分段整理 {chunk_elapsed}")
    if summary_elapsed:
        timing_parts.append(f"摘要生成 {summary_elapsed}")

    if timing_parts and total_elapsed:
        parts.append(f"- *耗時：{' ＋ '.join(timing_parts)}，總計約 {total_elapsed}。*")
    elif timing_parts:
        parts.append(f"- *耗時：{' ＋ '.join(timing_parts)}。*")
    elif total_elapsed:
        parts.append(f"- *耗時：總計約 {total_elapsed}。*")
    elif transcribe_sec > 0:
        parts.append(f"- *耗時：轉錄約 {int(transcribe_sec)} 秒。*")
    else:
        parts.append('- *耗時：未記錄詳細分段耗時。*')

    # ── 上傳歸檔 ──
    notion_label = 'Notion 課堂摘要庫' if metadata.get('type') == 'emba' else 'Notion 商務會談摘要 DB'
    parts.append(f"- *上傳歸檔：同步寫入 {service_emoji_tag('notion')} {notion_label} ＋ {service_emoji_tag('obsidian')} Obsidian 本地 Vault。*")

    # ── 免責聲明 ──
    parts.append('')
    parts.append('> ***⚠️ 本筆記摘要為 AI Agent :openclaw-dark:調用大語言模型（LLM）自動化 Workflow 生成。依據使用者提供錄音檔、加上個人筆記及其他補充材料進行自動交叉校對完成，但仍可能存在誤解、遺漏、錯別字或表述偏差；若需正式引用、對外使用或作為決策依據，請務必加以核實。***')

    # ── Obsidian 知識庫連結 ──
    parts.append('')
    parts.append('### 🗂 Obsidian 知識庫連結')
    if obsidian_md_path:
        rel_path = str(obsidian_md_path).split('小龍女知識庫/')[-1] if '小龍女知識庫/' in str(obsidian_md_path) else obsidian_md_path.name
        parts.append(f"- 📓 本堂課堂筆記　{rel_path}")
    course = metadata.get('course_name') or metadata.get('meeting_name') or ''
    if course:
        parts.append(f"- 📚 課程主檔　[[{course}]]")
    parts.append('- 📅 學期總覽　[[114-2 課程總覽]]')

    return "\n".join(parts).strip() + "\n"


def _build_obsidian_frontmatter(metadata: dict) -> str:
    def _esc(v: str) -> str:
        return (v or '').replace('"', "'")

    course = metadata.get('course_name') or metadata.get('meeting_name') or ''
    professor = metadata.get('professor') or ''
    room = metadata.get('room') or metadata.get('location') or ''
    keywords = metadata.get('keywords') or []
    if isinstance(keywords, str):
        keywords = [k.strip() for k in re.split(r'[,，、]+', keywords) if k.strip()]
    tags = ['lecture-transcribe', 'emba' if metadata.get('type') == 'emba' else 'business']
    lines = [
        '---',
        f"date: {metadata.get('date','')}",
        f'title: "{_esc(course)}"',
        f'professor: "{_esc(professor)}"',
        f'room: "{_esc(room)}"',
        f"type: {metadata.get('type','')}",
        'tags:',
    ]
    for t in tags + keywords:
        lines.append(f"  - {t}")
    lines.append('---\n')
    return "\n".join(lines)


def _should_include_transcript(metadata: dict) -> bool:
    """所有類型預設都不附完整逐字稿。"""
    return False


def _normalize_obsidian_markdown(md: str) -> str:
    """Post-process markdown for Obsidian compatibility.

    Two-pass approach:
      Pass 1 — Remove blank lines sandwiched between two table rows.
               (Prevents stray LLM-inserted blank lines from breaking table rendering.)
      Pass 2 — Ensure exactly one blank line before the first table row of a block,
               and at least one blank line after the last table row.
               (Obsidian requires a preceding blank line for tables to render.)

    This operates only on the Obsidian copy; the original summary_md fed to Notion
    is never touched, so Notion output is unaffected.
    """
    lines = md.split('\n')

    def _is_table_row(l: str) -> bool:
        return l.lstrip().startswith('|')

    # --- Pass 1: collapse blank lines between table rows ---
    pass1: list[str] = []
    for i, line in enumerate(lines):
        if line.strip() == '':
            prev_nb = next((lines[j] for j in range(i - 1, -1, -1) if lines[j].strip()), '')
            next_nb = next((lines[j] for j in range(i + 1, len(lines)) if lines[j].strip()), '')
            if _is_table_row(prev_nb) and _is_table_row(next_nb):
                continue  # drop spurious blank between rows
        pass1.append(line)

    # --- Pass 2: ensure blank lines at table block boundaries ---
    pass2: list[str] = []
    prev_is_table = False
    for line in pass1:
        cur_is_table = _is_table_row(line)
        if cur_is_table and not prev_is_table:
            # Start of table block — ensure preceding blank line
            if pass2 and pass2[-1].strip() != '':
                pass2.append('')
        elif not cur_is_table and prev_is_table and line.strip() != '':
            # End of table block — ensure following blank line
            if pass2 and pass2[-1].strip() != '':
                pass2.append('')
        pass2.append(line)
        prev_is_table = cur_is_table

    return '\n'.join(pass2)


def _write_obsidian_markdown(metadata: dict, summary_md: str, transcript_md: str) -> Optional[Path]:
    vault_path = _resolve_obsidian_vault_path()
    note_path = _obsidian_note_path(metadata, vault_path)
    if not note_path:
        return None
    note_path.parent.mkdir(parents=True, exist_ok=True)
    # Apply Obsidian-specific normalization (blank lines around tables, etc.)
    # This is separate from summary_md so Notion upload is never affected.
    obsidian_summary = _normalize_obsidian_markdown(summary_md.rstrip())
    content = _build_obsidian_frontmatter(metadata) + obsidian_summary
    if _should_include_transcript(metadata) and (transcript_md or '').strip():
        content += "\n\n---\n\n" + transcript_md.strip()
    content += "\n"
    note_path.write_text(content, encoding='utf-8')
    return note_path


# ============================================================
# 連結音訊下載（Google Drive / OneDrive 公開分享）
# ============================================================

_URL_RE = re.compile(r"https?://[^\s<>]+", re.I)
_AUDIO_EXTS = ('.mp3', '.m4a', '.wav', '.aac', '.flac', '.ogg', '.opus', '.wma', '.mp4', '.mov', '.webm')
_IMAGE_EXTS = ('.png', '.jpg', '.jpeg', '.webp', '.gif', '.heic')
_REFERENCE_EXTS = ('.pdf', '.pptx', '.docx', '.xlsx', '.xlsm', '.csv', '.tsv', '.md', '.markdown', '.txt', '.rtf')


def _extract_urls(text: str) -> list:
    if not (text or '').strip():
        return []
    urls = []
    seen = set()
    for m in _URL_RE.finditer(text):
        u = m.group(0).rstrip(').,]')
        if u in seen:
            continue
        seen.add(u)
        urls.append(u)
    return urls


def _iter_urls_with_line_context(text: str) -> list:
    items = []
    if not (text or '').strip():
        return items
    seen = set()
    for line in text.splitlines():
        for m in _URL_RE.finditer(line):
            u = m.group(0).rstrip(').,]')
            if u in seen:
                continue
            seen.add(u)
            items.append({'url': u, 'line': line.strip()})
    return items


def _is_notion_url(url: str) -> bool:
    u = (url or '').lower()
    return 'notion.so' in u


def _looks_like_reference_doc_url(url: str) -> bool:
    u = (url or '').lower()
    if _is_notion_url(u):
        return False
    return any(u.endswith(ext) for ext in _REFERENCE_EXTS)


def _looks_like_audio_url(url: str) -> bool:
    u = (url or '').lower()
    if _is_notion_url(u):
        return False
    if any(u.endswith(ext) for ext in _AUDIO_EXTS):
        return True
    # Common cloud share links are often audio in this workflow when not otherwise typed.
    if any(host in u for host in ('drive.google.com', 'docs.google.com', '1drv.ms', 'onedrive.live.com')) and not _looks_like_reference_doc_url(u):
        return True
    return False


def _want_diarization(text: str) -> bool:
    """Heuristic: only run diarization when the user explicitly asks for speaker separation.

    Default is OFF for stability/speed on long recordings.
    """
    t = (text or '').lower()
    keys = [
        'speaker', 'speakers', 'diarization',
        '分辨speaker', '分辨 speaker', '辨識speaker', '辨識 speaker',
        '說話者', '講者', '誰講', '誰在講',
    ]
    return any(k in t for k in keys)


def _pick_audio_url(text: str) -> str:
    items = _iter_urls_with_line_context(text)
    if not items:
        return ''

    audio_hints = ('錄音', '音檔', '音訊', '聲音', 'audio')
    for item in items:
        line = (item.get('line') or '').lower()
        url = item['url']
        if any(h in line for h in audio_hints) and _looks_like_audio_url(url):
            return url

    for item in items:
        url = item['url']
        if _looks_like_audio_url(url) and not _looks_like_reference_doc_url(url):
            return url

    for item in items:
        url = item['url']
        if not _is_notion_url(url) and not _looks_like_reference_doc_url(url):
            return url

    return items[0]['url']


def _pick_reference_urls(text: str) -> list:
    items = _iter_urls_with_line_context(text)
    if not items:
        return []

    audio_url = _pick_audio_url(text)
    out = []
    seen = set()
    ref_hints = ('筆記', 'notion', '教材', '講義', '附件', 'slides', 'slide', 'pdf', 'ppt', 'pptx', 'doc', 'docx', 'xls', 'xlsx')

    for item in items:
        url = item['url']
        line = (item.get('line') or '').lower()
        if url == audio_url:
            continue
        if _is_notion_url(url) or _looks_like_reference_doc_url(url) or any(h in line for h in ref_hints):
            if url not in seen:
                seen.add(url)
                out.append(url)

    return out


def _pick_first_url(text: str) -> str:
    return _pick_audio_url(text) or (_extract_urls(text)[0] if _extract_urls(text) else '')


def _extract_drive_file_id(url: str) -> str:
    """Extract Google Drive file id from common share links."""
    if not url:
        return ''
    # /file/d/<id>/
    m = re.search(r"/file/d/([a-zA-Z0-9_-]+)", url)
    if m:
        return m.group(1)
    # ?id=<id>
    m = re.search(r"[?&]id=([a-zA-Z0-9_-]+)", url)
    if m:
        return m.group(1)
    return ''


def _normalize_download_url(url: str) -> Tuple[str, str]:
    """Return (download_url, source_kind)."""
    u = (url or '').strip()
    if not u:
        return '', ''

    low = u.lower()

    # Google Drive
    if 'drive.google.com' in low or 'docs.google.com' in low:
        fid = _extract_drive_file_id(u)
        if fid:
            return f"https://drive.google.com/uc?export=download&id={fid}", 'gdrive'
        return u, 'gdrive'

    # OneDrive
    if '1drv.ms' in low or 'onedrive.live.com' in low:
        # Most share links accept download=1
        if 'download=1' not in low:
            sep = '&' if '?' in u else '?'
            u = f"{u}{sep}download=1"
        return u, 'onedrive'

    return u, 'generic'


def _http_download(url: str, out_path: Path, kind: str = 'generic') -> Path:
    """Download URL to out_path. Handles Google Drive confirm token for large files."""
    import urllib.request
    import urllib.parse

    out_path.parent.mkdir(parents=True, exist_ok=True)

    # cookie jar helps with Google Drive confirm token
    cj = None
    opener = None
    try:
        import http.cookiejar
        cj = http.cookiejar.CookieJar()
        opener = urllib.request.build_opener(urllib.request.HTTPCookieProcessor(cj))
    except Exception:
        opener = urllib.request.build_opener()

    def fetch(u: str):
        req = urllib.request.Request(u, headers={
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120 Safari/537.36'
        })
        return opener.open(req, timeout=60)

    # First request
    resp = fetch(url)

    # Google Drive sometimes returns an HTML page with a confirm token
    if kind == 'gdrive':
        ctype = (resp.headers.get('Content-Type') or '').lower()
        if 'text/html' in ctype:
            body = resp.read(200000).decode('utf-8', errors='ignore')
            mm = re.search(r"confirm=([0-9A-Za-z_]+)", body)
            if not mm:
                # older pattern
                mm = re.search(r"name=\"confirm\"\s+value=\"([0-9A-Za-z_]+)\"", body)
            if mm:
                token = mm.group(1)
                parsed = urllib.parse.urlparse(url)
                q = dict(urllib.parse.parse_qsl(parsed.query))
                q['confirm'] = token
                new_url = urllib.parse.urlunparse(parsed._replace(query=urllib.parse.urlencode(q)))
                resp = fetch(new_url)

    # Try infer filename
    cd = resp.headers.get('Content-Disposition') or ''
    fname = ''
    m = re.search(r'filename\*=UTF-8\'\'([^;]+)', cd)
    if m:
        fname = urllib.parse.unquote(m.group(1))
    else:
        m = re.search(r'filename="?([^";]+)"?', cd)
        if m:
            fname = m.group(1)

    final_path = out_path
    if fname:
        safe = fname.replace('/', '-')
        final_path = out_path.with_name(safe)

    with open(final_path, 'wb') as f:
        while True:
            chunk = resp.read(1024 * 1024)
            if not chunk:
                break
            f.write(chunk)

    return final_path


def _download_shared_audio(url: str, out_dir: Path) -> Path:
    dl, kind = _normalize_download_url(url)
    if not dl:
        raise ValueError('empty url')

    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    out = out_dir / f"link_{ts}.bin"
    p = _http_download(dl, out, kind=kind)
    return p


def _download_reference_file(url: str, out_dir: Path) -> Path:
    dl, kind = _normalize_download_url(url)
    if not dl:
        raise ValueError('empty url')

    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    out = out_dir / f"ref_{ts}.bin"
    return _http_download(dl, out, kind=kind)


def _notion_headers() -> dict:
    key = (os.environ.get('NOTION_API_KEY') or '').strip()
    if not key:
        return {}
    return {
        'Authorization': f'Bearer {key}',
        'Notion-Version': '2025-09-03',
        'Content-Type': 'application/json',
    }


def _notion_get_json(url: str) -> dict:
    headers = _notion_headers()
    if not headers:
        raise RuntimeError('NOTION_API_KEY 未設定')

    import urllib.request
    req = urllib.request.Request(url, headers=headers)
    with urllib.request.urlopen(req, timeout=30) as resp:
        return json.loads(resp.read().decode('utf-8'))


def _extract_notion_page_id(url: str) -> str:
    """Extract the 32-hex-char page id from a Notion share URL.

    Strategy: split the path slug by '-' and find the last 32-char hex segment.
    This avoids false matches caused by stripping dashes from the entire URL.
    """
    if not url:
        return ''
    path = url.split('?')[0].split('#')[0]
    slug = path.rstrip('/').rsplit('/', 1)[-1]
    parts = slug.split('-')
    for part in reversed(parts):
        if re.fullmatch(r'[0-9a-fA-F]{32}', part):
            raw = part
            return f"{raw[0:8]}-{raw[8:12]}-{raw[12:16]}-{raw[16:20]}-{raw[20:32]}"
    # fallback: try finding 8-4-4-4-12 uuid pattern in full url
    m = re.search(r'([0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12})', url)
    if m:
        return m.group(1)
    return ''


def _notion_rich_text_plain(rich_text) -> str:
    parts = []
    for item in rich_text or []:
        if not isinstance(item, dict):
            continue
        plain = item.get('plain_text')
        if plain is None:
            plain = (((item.get('text') or {}).get('content')) or '')
        if plain:
            parts.append(str(plain))
    return ''.join(parts).strip()


def _notion_page_title(page: dict) -> str:
    props = (page or {}).get('properties') or {}
    for prop in props.values():
        if not isinstance(prop, dict):
            continue
        if prop.get('type') == 'title':
            title = _notion_rich_text_plain(prop.get('title') or [])
            if title:
                return title
    return ''


def _notion_block_file_url(block: dict) -> str:
    if not isinstance(block, dict):
        return ''
    typ = block.get('type') or ''
    data = block.get(typ) or {}
    if typ not in {'image', 'file'}:
        return ''
    if isinstance(data, dict):
        # Notion commonly returns {type: {file: {url}}, but some payloads
        # surface {type: {url}} directly. Support both.
        if isinstance(data.get('file'), dict):
            return (data.get('file') or {}).get('url', '').strip()
        if isinstance(data.get('external'), dict):
            return (data.get('external') or {}).get('url', '').strip()
        return (data.get('url') or '').strip()
    return ''


def _notion_block_rich_text(block: dict) -> str:
    if not isinstance(block, dict):
        return ''
    typ = block.get('type') or ''
    data = block.get(typ) or {}
    if typ in {'paragraph', 'heading_1', 'heading_2', 'heading_3', 'bulleted_list_item', 'numbered_list_item', 'to_do', 'toggle', 'quote', 'callout', 'code'}:
        return _notion_rich_text_plain(data.get('rich_text') or [])
    if typ == 'table_row':
        cells = []
        for cell in data.get('cells') or []:
            cells.append(_notion_rich_text_plain(cell))
        return ' | '.join([c for c in cells if c])
    if typ in {'child_page', 'child_database'}:
        return data.get('title') or ''
    if typ == 'bookmark':
        return data.get('caption') or data.get('url') or ''
    if typ == 'equation':
        return (data.get('expression') or '').strip()
    if typ in {'image', 'file'}:
        caption = _notion_rich_text_plain(data.get('caption') or [])
        url = _notion_block_file_url(block)
        if caption and url:
            return f'{caption} ({url})'
        return caption or url
    if typ == 'transcription':
        parts = []
        for key, label in (('summary', 'AI 摘要'), ('transcript', '逐字稿'), ('rich_text', '內容')):
            val = data.get(key)
            if isinstance(val, list):
                txt = _notion_rich_text_plain(val)
            else:
                txt = (val or '').strip() if isinstance(val, str) else ''
            if txt:
                parts.append(f'{label}: {txt}')
        return '\n'.join(parts).strip()
    return ''


def _notion_block_to_text(block: dict, depth: int = 0) -> str:
    if not isinstance(block, dict):
        return ''
    typ = block.get('type') or ''
    data = block.get(typ) or {}
    indent = '  ' * max(depth, 0)
    rich = _notion_block_rich_text(block)

    rich = (rich or '').strip()
    if not rich:
        return ''

    if typ == 'heading_1':
        return f"# {rich}"
    if typ == 'heading_2':
        return f"## {rich}"
    if typ == 'heading_3':
        return f"### {rich}"
    if typ == 'bulleted_list_item':
        return f"{indent}- {rich}"
    if typ == 'numbered_list_item':
        return f"{indent}1. {rich}"
    if typ == 'to_do':
        checked = 'x' if data.get('checked') else ' '
        return f"{indent}- [{checked}] {rich}"
    if typ == 'quote':
        return f"> {rich}"
    if typ == 'callout':
        return f"{indent}💡 {rich}"
    if typ == 'code':
        return f"{indent}```\n{rich}\n```"
    if typ in {'image', 'file'}:
        return f"{indent}[{typ}] {rich}"
    return f"{indent}{rich}"


def _collect_notion_media_blocks(block_id: str, depth: int = 0, max_depth: int = 2, max_blocks: int = 200) -> list:
    items = []
    cursor = None
    seen = 0

    while True:
        url = f'https://api.notion.com/v1/blocks/{block_id}/children?page_size=100'
        if cursor:
            url += f'&start_cursor={cursor}'
        data = _notion_get_json(url)
        for block in data.get('results') or []:
            if seen >= max_blocks:
                return items
            seen += 1
            if not isinstance(block, dict):
                continue
            typ = block.get('type') or ''
            if typ in {'image', 'file'}:
                src_url = _notion_block_file_url(block)
                if src_url:
                    media = block.get(typ) or {}
                    caption = _notion_rich_text_plain(media.get('caption') or [])
                    filename_hint = ''
                    if typ == 'file' and isinstance(media.get('file'), dict):
                        filename_hint = (media.get('file') or {}).get('name') or ''
                    items.append({
                        'type': typ,
                        'url': src_url,
                        'caption': caption,
                        'filename_hint': filename_hint,
                        'id': block.get('id') or '',
                    })
            if block.get('has_children') and depth < max_depth:
                items.extend(_collect_notion_media_blocks(block.get('id') or '', depth=depth + 1, max_depth=max_depth, max_blocks=max(20, max_blocks - seen)))
        if not data.get('has_more'):
            break
        cursor = data.get('next_cursor')
        if not cursor:
            break
    return items


def _fetch_notion_children_text(block_id: str, depth: int = 0, max_depth: int = 2, max_blocks: int = 200) -> list:
    lines = []
    cursor = None
    seen = 0

    while True:
        url = f'https://api.notion.com/v1/blocks/{block_id}/children?page_size=100'
        if cursor:
            url += f'&start_cursor={cursor}'
        data = _notion_get_json(url)
        for block in data.get('results') or []:
            if seen >= max_blocks:
                return lines
            seen += 1
            line = _notion_block_to_text(block, depth=depth)
            if line:
                lines.append(line)
            if block.get('has_children') and depth < max_depth:
                child_lines = _fetch_notion_children_text(block.get('id') or '', depth=depth + 1, max_depth=max_depth, max_blocks=max(20, max_blocks - seen))
                lines.extend(child_lines)
        if not data.get('has_more'):
            break
        cursor = data.get('next_cursor')
        if not cursor:
            break
    return lines


def _extract_notion_transcription_sections(block_id: str, depth: int = 0, max_depth: int = 2, max_blocks: int = 200) -> list:
    items = []
    cursor = None
    seen = 0

    while True:
        url = f'https://api.notion.com/v1/blocks/{block_id}/children?page_size=100'
        if cursor:
            url += f'&start_cursor={cursor}'
        data = _notion_get_json(url)
        for block in data.get('results') or []:
            if seen >= max_blocks:
                return items
            seen += 1
            if not isinstance(block, dict):
                continue
            if (block.get('type') or '') == 'transcription':
                payload = block.get('transcription') or {}
                items.append({
                    'summary': (payload.get('summary') or '').strip() if isinstance(payload.get('summary'), str) else _notion_rich_text_plain(payload.get('summary') or []),
                    'transcript': (payload.get('transcript') or '').strip() if isinstance(payload.get('transcript'), str) else _notion_rich_text_plain(payload.get('transcript') or []),
                    'content': _notion_block_rich_text(block),
                })
            if block.get('has_children') and depth < max_depth:
                items.extend(_extract_notion_transcription_sections(block.get('id') or '', depth=depth + 1, max_depth=max_depth, max_blocks=max(20, max_blocks - seen)))
        if not data.get('has_more'):
            break
        cursor = data.get('next_cursor')
        if not cursor:
            break
    return items


def _fetch_notion_reference_text(url: str, max_chars: Optional[int] = None) -> Tuple[str, dict]:
    page_id = _extract_notion_page_id(url)
    if not page_id:
        raise ValueError('無法從 Notion 連結解析 page id')

    page = _notion_get_json(f'https://api.notion.com/v1/pages/{page_id}')
    title = _notion_page_title(page) or 'Notion 筆記'
    lines = [f"# {title}"]
    lines.extend(_fetch_notion_children_text(page_id, depth=0, max_depth=2, max_blocks=240))
    text = "\n".join([x for x in lines if (x or '').strip()]).strip()
    limit = max_chars or CONFIG.get('max_reference_chars_per_source') or 6000
    if len(text) > limit:
        text = text[:limit].rstrip() + "\n...(節錄)"
    media = _collect_notion_media_blocks(page_id, depth=0, max_depth=2, max_blocks=240)
    transcription_sections = _extract_notion_transcription_sections(page_id, depth=0, max_depth=2, max_blocks=240)
    return text, {'title': title, 'page_id': page_id, 'url': url, 'media_blocks': media, 'transcription_sections': transcription_sections}


# ============================================================
# 課表推斷
# ============================================================

def _load_schedule():
    p = CONFIG["schedule_path"]
    if not p.exists():
        return []
    with open(p, "r", encoding="utf-8") as f:
        return yaml.safe_load(f).get("courses", [])


def _normalize_text_for_match(text: str) -> str:
    return re.sub(r'\s+', '', (text or '').strip()).lower()


def _match_schedule_course(course_name: str) -> Optional[dict]:
    norm = _normalize_text_for_match(course_name)
    if not norm:
        return None
    for course in _load_schedule():
        name = course.get('name') or ''
        course_norm = _normalize_text_for_match(name)
        if norm == course_norm or norm in course_norm or course_norm in norm:
            return course
    return None


def _extract_explicit_course_from_text(text: str) -> str:
    t = (text or '').strip()
    if not t:
        return ''
    for course in _load_schedule():
        name = (course.get('name') or '').strip()
        if name and name in t:
            return name
    return ''


def _apply_user_context_overrides(metadata: dict, provided: dict, ctx_text: str) -> dict:
    provided = provided or {}

    explicit_course = ''
    for key in _EXPLICIT_COURSE_KEYS + _DEF_TOPIC_KEYS:
        val = (provided.get(key) or '').strip()
        if val:
            explicit_course = val
            break
    if not explicit_course:
        explicit_course = _extract_explicit_course_from_text(ctx_text)

    explicit_professor = ''
    for key in _EXPLICIT_PROFESSOR_KEYS:
        val = (provided.get(key) or '').strip()
        if val:
            explicit_professor = val
            break

    if explicit_course:
        metadata['course_name_explicit'] = explicit_course
        metadata['course_name'] = explicit_course
        metadata['meeting_name'] = explicit_course
        metadata['auto_detected'] = False

        matched = _match_schedule_course(explicit_course)
        if matched:
            metadata['course_name'] = matched.get('name') or explicit_course
            metadata['meeting_name'] = matched.get('name') or explicit_course
            metadata['professor'] = matched.get('professor') or metadata.get('professor') or ''
            metadata['room'] = matched.get('room') or metadata.get('room') or ''
            metadata['type'] = 'emba'
            metadata['schedule_matched_by_user_context'] = True

    if explicit_professor:
        metadata['professor'] = explicit_professor

    return metadata


def _infer_course(dt=None):
    if dt is None:
        dt = datetime.now()
    courses = _load_schedule()
    iso_wd = dt.weekday() + 1
    cur_time = dt.strftime("%H:%M")

    for c in courses:
        if c["day"] == iso_wd and c["start"] <= cur_time <= c["end"]:
            return {
                "course_name": c["name"], "professor": c["professor"],
                "room": c.get("room", ""), "date": dt.strftime("%Y-%m-%d"),
                "day_of_week": "一二三四五六日"[dt.weekday()],
                "auto_detected": True, "type": "emba",
            }
    return {
        "course_name": "", "professor": "", "room": "",
        "date": dt.strftime("%Y-%m-%d"),
        "day_of_week": "一二三四五六日"[dt.weekday()],
        "auto_detected": False, "type": "",
    }



# ============================================================
# 互動式補齊使用者提供的會議/課堂資訊
# ============================================================

_DEF_TOPIC_KEYS = ["摘要主題", "主題", "課程", "會議"]
_EXPLICIT_COURSE_KEYS = ["課程", "課名", "課程名稱", "class", "course"]
_EXPLICIT_PROFESSOR_KEYS = ["教授", "老師", "授課老師", "講師"]


def _extract_reference_entities(text: str) -> list:
    """Extract likely people / company / topic entities from user-provided notes.

    Keep this intentionally lightweight to avoid token bloat.
    """
    t = (text or '').strip()
    if not t:
        return []

    cand = []
    for line in re.split(r'[\n,，、/｜|；;]+', t):
        s = line.strip().strip('-•*#：:')
        if not s:
            continue
        if len(s) < 2 or len(s) > 40:
            continue
        if re.search(r'https?://', s, re.I):
            continue
        if re.search(r'^(ok|跳過|模板|關鍵字|補充|摘要|筆記)$', s, re.I):
            continue
        if re.search(r'^(以下是|我的|個人)?[^\n]{0,12}筆記', s):
            continue
        # strong business / talk entities
        if re.search(r'(公司|科技|工業|電|大學|學院|教授|主任|經理|協理|董事長|總經理|副總裁|論壇|演講|課程|液冷|散熱|BBU|UPS|BMS|CDU|manifold|Quick Disconnect)', s, re.I):
            cand.append(s)
            continue
        # probable Chinese names / organizations / short technical terms
        if re.fullmatch(r'[\u4e00-\u9fffA-Za-z0-9\-\+\.& ]{2,20}', s):
            cand.append(s)

    out = []
    seen = set()
    for s in cand:
        k = s.lower()
        if k in seen:
            continue
        seen.add(k)
        out.append(s)
        if len(out) >= 24:
            break
    return out


_COMPANY_STOPWORDS = {
    '老師', '教授', '同學', '學生', '市場', '產業', '策略', '個案', '案例', '課程', '課堂',
    '學校', '大學', '學院', '管理', '研究', '董事長', '總經理', '執行長', '台灣', '中國', '美國',
    '日本', '越南', '泰國', '印尼', '亞洲', '歐洲', '老師補充', '使用者補充筆記',
}


def _normalize_company_candidate_name(text: str) -> str:
    s = re.sub(r'^[\-•*#\d\.\)\(\[\]\s:：]+', '', (text or '').strip())
    s = re.sub(r'^(今天談到|今天談|老師提到|老師講到|教材列出|教材提到|案例是|案例談到|例如|像是|包括|以及|並比較)', '', s)
    s = re.sub(r'[（(].{0,18}?[)）]$', '', s).strip()
    s = re.sub(r'\s+', ' ', s)
    return s[:60]


def _extract_company_candidates_from_text(text: str, source: str = 'unknown') -> list:
    t = (text or '').strip()
    if not t:
        return []

    patterns = [
        r'([\u4e00-\u9fffA-Za-z0-9&\-·\.]{2,40}(?:公司|集團|科技|工業|實業|企業|控股|電腦|電子|半導體|通訊|食品|銀行|金控|航空|海運|零售|製造))',
        r'\b([A-Z][A-Za-z0-9&\-\.]{1,28}(?:\s+[A-Z][A-Za-z0-9&\-\.]{1,28}){0,3})\b',
    ]

    out = []
    seen = set()
    for pat in patterns:
        for raw in re.findall(pat, t):
            name = _normalize_company_candidate_name(raw)
            if not name or len(name) < 2:
                continue
            key = name.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append({'name': name, 'source': source})

    if source in {'reference_entities', 'reference_notes', 'image_ocr_notes'}:
        for token in re.split(r'[\n,，、/｜|；;：: ]+', t):
            name = _normalize_company_candidate_name(token)
            if not name or len(name) < 2 or len(name) > 12:
                continue
            if not re.fullmatch(r'[\u4e00-\u9fffA-Za-z0-9&\-\.]{2,12}', name):
                continue
            key = name.lower()
            if key in seen:
                continue
            seen.add(key)
            out.append({'name': name, 'source': source})
    return out


def _score_company_candidate(name: str, source_hits: list[str], transcript_text: str) -> int:
    score = 0
    if any(token in name for token in ('公司', '集團', '科技', '工業', '實業', '企業')):
        score += 3
    if any(src in ('course_name', 'reference_notes', 'ocr_notes', 'image_ocr_notes') for src in source_hits):
        score += 2
    mention_count = transcript_text.count(name)
    score += min(mention_count, 4)
    case_context = 0
    for kw in ('個案', '案例', '策略', '市場', '轉型', '風險', '競爭', '決策'):
        if kw in transcript_text:
            case_context += 1
    if case_context:
        score += 1
    return score


def _select_company_candidates(metadata: dict, transcript_text: str, top_n: int = 5) -> list:
    sources = []
    for key in ('course_name', 'reference_notes', 'image_ocr_notes'):
        val = metadata.get(key) or ''
        if val:
            sources.append((key, val))
    ref_entities = metadata.get('reference_entities') or []
    if ref_entities:
        sources.append(('reference_entities', '\n'.join([str(x) for x in ref_entities if x])))
    if transcript_text:
        sources.append(('transcript', transcript_text))

    bucket = {}
    for source_name, text in sources:
        for item in _extract_company_candidates_from_text(text, source=source_name):
            name = item['name']
            key = name.lower()
            if key in {x.lower() for x in _COMPANY_STOPWORDS}:
                continue
            if any(bad in name for bad in ('老師', '教授', '課程', '個案', '案例', '市場的策略')):
                continue
            if len(name) <= 2 and not re.search(r'[A-Z]{2,}', name):
                continue
            row = bucket.setdefault(key, {'company': name, 'sources': set()})
            row['sources'].add(source_name)

    ranked = []
    for row in bucket.values():
        company = row['company']
        if company in _COMPANY_STOPWORDS:
            continue
        score = _score_company_candidate(company, list(row['sources']), transcript_text)
        if score <= 0:
            continue
        ranked.append({
            'company': company,
            'score': score,
            'sources': sorted(row['sources']),
            'mention_count': transcript_text.count(company),
        })

    ranked.sort(key=lambda x: (-x['score'], -x['mention_count'], x['company']))
    return ranked[:top_n]


def _build_company_enrichment_plan(metadata: dict, transcript_text: str, top_n: int = 5) -> list:
    candidates = _select_company_candidates(metadata, transcript_text, top_n=top_n)
    plan = []
    for item in candidates:
        company = item['company']
        plan.append({
            **item,
            'query': f'{company} 公司背景 產業 市場 台商 個案',
            'status': 'pending_external_search',
        })
    return plan


def _prepare_company_web_enrichment(metadata: dict, transcript_text: str) -> list:
    if not metadata.get('enable_company_web_enrichment'):
        return []
    existing = metadata.get('company_web_enrichment') or []
    if existing:
        return existing
    plan = _build_company_enrichment_plan(metadata, transcript_text, top_n=5)
    metadata['company_candidates'] = plan
    metadata['company_web_enrichment'] = [
        {
            'company': item['company'],
            'query': item['query'],
            'summary': (
                '尚未實際查詢。此欄位為 web enrichment slot，後續可由 OpenClaw web_search 結果填入，'
                '再交給 final model 整理成「外部補充資訊」。'
            ),
            'status': item['status'],
        }
        for item in plan
    ]
    return metadata['company_web_enrichment']


def _build_reference_entities(metadata: dict, ctx_text: str) -> list:
    items = []
    for key in ('course_name', 'meeting_name', 'company', 'target_company', 'professor'):
        v = (metadata.get(key) or '').strip()
        if v:
            items.extend(_extract_reference_entities(v))
    for key in ('attendees', 'participants', 'keywords'):
        vals = metadata.get(key) or []
        if isinstance(vals, str):
            vals = re.split(r'[,，、/｜|;；]+', vals)
        for v in vals:
            vv = (str(v).strip())
            if vv:
                items.extend(_extract_reference_entities(vv))
    notes = metadata.get('reference_notes') or ''
    if notes:
        items.extend(_extract_reference_entities(notes))
    if ctx_text:
        items.extend(_extract_reference_entities(ctx_text))

    out = []
    seen = set()
    for s in items:
        k = s.lower()
        if k in seen:
            continue
        seen.add(k)
        out.append(s)
        if len(out) >= 32:
            break
    return out


def _collect_attachment_candidates(payload: dict, **kwargs) -> list:
    candidates = []
    # NOTE: 'reference_file_paths' is explicitly included so that callers can pass
    # PDF/PPTX/DOCX paths via payload['reference_file_paths'] and have them routed
    # through _collect_reference_file_paths correctly.
    keys = (
        'attachments', 'files', 'media', 'images', 'documents', 'docs', 'document',
        'file', 'image_path', 'image', 'file_path', 'path',
        'reference_file_paths',  # ← FIX: was missing; PDFs passed via this key were silently dropped
    )

    if isinstance(payload, dict):
        for key in keys:
            v = payload.get(key)
            if v:
                candidates.append(v)
    for key in keys:
        v = kwargs.get(key)
        if v:
            candidates.append(v)

    flat = []
    for c in candidates:
        if isinstance(c, (list, tuple)):
            flat.extend(c)
        else:
            flat.append(c)
    return flat


def _item_to_path_and_mime(item) -> Tuple[str, str]:
    if isinstance(item, dict):
        p = item.get('path') or item.get('file_path') or item.get('localPath') or item.get('image_path')
        mime = (item.get('mimeType') or item.get('mime') or '').lower()
        return str(p) if p else '', mime
    if isinstance(item, str):
        return item, ''
    return '', ''


def _is_image_path(path: str, mime: str = '') -> bool:
    p = (path or '').lower()
    m = (mime or '').lower()
    return bool(p) and (m.startswith('image/') or p.endswith(_IMAGE_EXTS))


def _is_audio_path(path: str, mime: str = '') -> bool:
    p = (path or '').lower()
    m = (mime or '').lower()
    return bool(p) and (m.startswith('audio/') or p.endswith(_AUDIO_EXTS))


def _is_reference_material_path(path: str, mime: str = '') -> bool:
    p = (path or '').lower()
    m = (mime or '').lower()
    if not p:
        return False
    if _is_image_path(p, m) or _is_audio_path(p, m):
        return False
    return m in (
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'text/plain', 'text/markdown', 'text/csv', 'text/tab-separated-values',
    ) or p.endswith(_REFERENCE_EXTS)


def _dedupe_paths(paths: list, limit: int = 8) -> list:
    out = []
    seen = set()
    for p in paths:
        pp = str(p)
        if not pp or pp in seen:
            continue
        seen.add(pp)
        out.append(pp)
        if len(out) >= limit:
            break
    return out


def _collect_image_paths(payload: dict, **kwargs) -> list:
    paths = []
    for item in _collect_attachment_candidates(payload, **kwargs):
        p, mime = _item_to_path_and_mime(item)
        if _is_image_path(p, mime):
            paths.append(p)
    return _dedupe_paths(paths, limit=8)


def _collect_reference_file_paths(payload: dict, **kwargs) -> list:
    paths = []
    for item in _collect_attachment_candidates(payload, **kwargs):
        p, mime = _item_to_path_and_mime(item)
        if _is_reference_material_path(p, mime):
            paths.append(p)
    return _dedupe_paths(paths, limit=10)


def _ocr_image_texts(image_paths: list) -> list:
    if not image_paths:
        return []
    script = Path(__file__).parent / 'scripts' / 'ocr_image.swift'
    if not script.exists():
        return []
    try:
        res = subprocess.run(
            ['swift', str(script), *image_paths],
            capture_output=True,
            text=True,
            timeout=45,
            check=True,
        )
        data = json.loads(res.stdout or '[]')
        return [x for x in data if isinstance(x, dict)]
    except Exception:
        return []


def _record_ocr_stats(metadata: dict, source_label: str, image_paths: list, ocr_rows: list) -> dict:
    proc = metadata.setdefault('process_summary', {})
    bucket = proc.setdefault('ocr_sources', [])
    success_paths = set()
    total_chars = 0
    success_count = 0
    for row in ocr_rows or []:
        if not isinstance(row, dict):
            continue
        text = (row.get('text') or row.get('ocr_text') or row.get('content') or '').strip()
        src = str(row.get('path') or row.get('image_path') or '')
        if text:
            success_count += 1
            total_chars += len(text)
            if src:
                success_paths.add(src)
    stats = {
        'source': source_label,
        'image_count': len(image_paths or []),
        'success_count': success_count,
        'failed_count': max(len(image_paths or []) - len(success_paths), 0),
        'total_chars': total_chars,
        'warning': '',
    }
    if stats['image_count'] and stats['total_chars'] == 0:
        stats['warning'] = 'OCR 未抽到可用文字'
    bucket.append(stats)
    proc['ocr_summary'] = {
        'image_count': sum(x.get('image_count', 0) for x in bucket),
        'success_count': sum(x.get('success_count', 0) for x in bucket),
        'failed_count': sum(x.get('failed_count', 0) for x in bucket),
        'total_chars': sum(x.get('total_chars', 0) for x in bucket),
        'warning': '；'.join([x.get('warning', '') for x in bucket if x.get('warning')]).strip(),
    }
    return stats


def _truncate_reference_text(text: str, max_chars: Optional[int] = None) -> str:
    t = (text or '').strip()
    if not t:
        return ''
    limit = max_chars or CONFIG.get('max_reference_chars_per_source') or 6000
    if len(t) <= limit:
        return t
    return t[:limit].rstrip() + "\n...(節錄)"


def _truncate_ocr_reference_text(text: str, max_chars: Optional[int] = None) -> str:
    t = (text or '').strip()
    if not t:
        return ''
    limit = max_chars or CONFIG.get('max_ocr_reference_chars_per_source') or 9000
    if len(t) <= limit:
        return t
    return t[:limit].rstrip() + "\n...(教材 OCR 節錄，保留前段高密度內容)"


def _download_notion_media(url: str, out_dir: Path, hint: str = 'media') -> Path:
    dl, kind = _normalize_download_url(url)
    if not dl:
        raise ValueError('empty url')
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    ext = Path(dl.split('?', 1)[0]).suffix[:12]
    safe_ext = ext if re.fullmatch(r'\.[A-Za-z0-9]{1,10}', ext or '') else ''
    out = out_dir / f'notion_{hint}_{ts}{safe_ext or ".bin"}'
    return _http_download(dl, out, kind=kind)


def _classify_notion_media_item(item: dict) -> str:
    typ = (item.get('type') or '').lower()
    url = (item.get('url') or '').strip()
    hint = (item.get('filename_hint') or '').strip().lower()
    pathish = f"{url.split('?', 1)[0]} {hint}".strip()
    if typ == 'image':
        return 'image'
    if typ == 'file' and re.search(r'\.(png|jpe?g|gif|webp|bmp|tiff?)$', pathish, re.I):
        return 'image'
    if re.search(r'\.(pdf|docx|pptx|xlsx|xlsm|csv|tsv|txt|md|markdown|rtf)$', pathish, re.I):
        return 'document'
    return 'unknown'


def _maybe_ocr_reference_images(image_paths: list, source_title: str) -> str:
    if not image_paths:
        return ''
    ocr_rows = _ocr_image_texts(image_paths)
    if not ocr_rows:
        return ''
    chunks = [
        f'## 教材照片 OCR 詳細還原：{source_title}',
        '- 用途：這一段是教材翻拍照片 / 投影片 / 講義內容的 OCR 重建素材，請優先用來還原定義、分類、步驟、比較點、判準、條列與考點。',
        '- 規則：若內容明顯屬於考試重點或概念理解核心，寧可保留較多原始細節，也不要過度濃縮成泛泛幾句。',
        '',
    ]
    success_count = 0
    for idx, row in enumerate(ocr_rows, 1):
        if not isinstance(row, dict):
            continue
        text = _truncate_ocr_reference_text(row.get('text') or row.get('ocr_text') or row.get('content') or '')
        if not text:
            continue
        success_count += 1
        src = row.get('path') or row.get('image_path') or ''
        name = Path(str(src)).name if src else f'image_{idx}'
        chunks.append(f'### 圖 {idx}｜{name}')
        chunks.append('- 請盡量按原文結構還原，不要先替它做過度摘要。')
        chunks.append(text)
        chunks.append('')
    if success_count == 0:
        return ''
    return '\n'.join(chunks).strip()


def _append_reference_section(metadata: dict, title: str, body: str) -> bool:
    is_ocr_material = any(key in title for key in ('OCR', '教材照片', '圖片/OCR'))
    truncator = _truncate_ocr_reference_text if is_ocr_material else _truncate_reference_text
    body = truncator(body)
    if not body:
        return False

    section = f'[{title}]\n{body}'
    existing = (metadata.get('reference_notes') or '').strip()
    if section in existing:
        return False

    total_limit = CONFIG.get('max_reference_chars_total') or 18000
    if existing:
        sep = '\n\n'
        remain = total_limit - len(existing) - len(sep) - len(f'[{title}]\n')
        if remain <= 80:
            return False
        body = truncator(body, max_chars=remain)
        section = f'[{title}]\n{body}'
        metadata['reference_notes'] = existing + sep + section
    else:
        metadata['reference_notes'] = truncator(section, max_chars=total_limit)

    sources = metadata.setdefault('reference_sources', [])
    if title not in sources:
        sources.append(title)
    if is_ocr_material:
        metadata['detailed_material_restore'] = True
        metadata['high_priority_exam_material'] = True
    return True


def _store_notion_ai_transcript_candidate(metadata: dict, title: str, body: str) -> bool:
    body = _truncate_reference_text(body)
    if not body:
        return False

    candidates = metadata.setdefault('notion_ai_transcript_candidates', [])
    key = f'{title}\n{body[:160]}'
    for item in candidates:
        if item.get('key') == key:
            return False

    candidates.append({
        'title': title,
        'body': body,
        'char_len': len(body),
        'key': key,
    })
    source_title = f'Notion AI 逐字稿（保留備援）：{title}'
    sources = metadata.setdefault('reference_sources', [])
    if source_title not in sources:
        sources.append(source_title)
    metadata['has_notion_ai_transcript_candidate'] = True
    return True


def _assess_transcript_quality(segments: list, transcript_text: str, duration_sec: float) -> dict:
    rows = []
    for seg in segments or []:
        text = (seg.get('text') or '').strip()
        if text:
            rows.append(text)

    duration_min = max((duration_sec or 0) / 60, 1.0)
    text_len = len((transcript_text or '').strip())
    chars_per_min = text_len / duration_min if duration_min > 0 else 0.0
    avg_seg_chars = (sum(len(x) for x in rows) / len(rows)) if rows else 0.0
    short_seg_ratio = (sum(1 for x in rows if len(x) <= 4) / len(rows)) if rows else 1.0

    normalized = []
    for row in rows:
        norm = re.sub(r'\s+', '', row.lower())
        if norm:
            normalized.append(norm)
    unique_ratio = (len(set(normalized)) / len(normalized)) if normalized else 0.0
    repeated_seg_ratio = 1.0 - unique_ratio if normalized else 1.0

    score = 0
    reasons = []
    if chars_per_min < 90:
        score += 2
        reasons.append(f'逐字稿密度偏低（約 {chars_per_min:.0f} 字/分）')
    elif chars_per_min < 130:
        score += 1
        reasons.append(f'逐字稿密度略低（約 {chars_per_min:.0f} 字/分）')

    if avg_seg_chars < 8:
        score += 1
        reasons.append(f'平均每段字數偏短（約 {avg_seg_chars:.1f} 字）')

    if short_seg_ratio >= 0.45:
        score += 1
        reasons.append(f'極短片段比例偏高（{short_seg_ratio:.0%}）')

    if repeated_seg_ratio >= 0.35:
        score += 1
        reasons.append(f'重複片段比例偏高（{repeated_seg_ratio:.0%}）')

    poor = score >= 2
    if poor:
        label = 'poor'
    elif score == 1:
        label = 'fair'
    else:
        label = 'good'

    return {
        'label': label,
        'poor': poor,
        'score': score,
        'text_len': text_len,
        'chars_per_min': round(chars_per_min, 1),
        'avg_seg_chars': round(avg_seg_chars, 1),
        'short_seg_ratio': round(short_seg_ratio, 3),
        'repeated_seg_ratio': round(repeated_seg_ratio, 3),
        'reasons': reasons,
    }


def _maybe_enable_notion_ai_transcript_fallback(metadata: dict, quality: dict) -> bool:
    metadata['original_transcript_quality'] = quality.get('label') or 'unknown'
    metadata['original_transcript_quality_reasons'] = quality.get('reasons') or []

    if not metadata.get('skip_notion_ai_transcript'):
        return False

    candidates = metadata.get('notion_ai_transcript_candidates') or []
    if not candidates:
        return False

    metadata['notion_ai_transcript_reference_only'] = True
    if not quality.get('poor'):
        return False

    chosen = max(candidates, key=lambda item: item.get('char_len') or 0)
    title = (chosen.get('title') or 'Notion AI 逐字稿').strip()
    body = (chosen.get('body') or '').strip()
    if not body:
        return False

    ok = _append_reference_section(metadata, f'Notion AI 逐字稿（fallback 啟用）：{title}', body)
    if ok:
        metadata['use_notion_ai_transcript_fallback'] = True
        metadata['notion_ai_transcript_fallback_title'] = title
    return ok


def _read_plain_text_file(path: Path, max_chars: Optional[int] = None) -> str:
    try:
        return _truncate_reference_text(path.read_text(encoding='utf-8', errors='ignore'), max_chars=max_chars)
    except Exception:
        try:
            return _truncate_reference_text(path.read_text(encoding='utf-8-sig', errors='ignore'), max_chars=max_chars)
        except Exception:
            return ''


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        lines = []
        for p in doc.paragraphs:
            t = (p.text or '').strip()
            if t:
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                vals = [(c.text or '').strip() for c in row.cells]
                vals = [v for v in vals if v]
                if vals:
                    lines.append(' | '.join(vals))
        return _truncate_reference_text('\n'.join(lines))
    except Exception:
        try:
            return _truncate_reference_text(subprocess.check_output(['textutil', '-convert', 'txt', '-stdout', str(path)], text=True, stderr=subprocess.DEVNULL))
        except Exception:
            return ''


def _extract_pptx_text(path: Path) -> str:
    try:
        slides = []
        with zipfile.ZipFile(path) as zf:
            names = sorted([n for n in zf.namelist() if re.fullmatch(r'ppt/slides/slide\d+\.xml', n)])
            for idx, name in enumerate(names, 1):
                root = ET.fromstring(zf.read(name))
                texts = [t.text.strip() for t in root.iter() if t.tag.endswith('}t') and (t.text or '').strip()]
                if texts:
                    slides.append(f'Slide {idx}: ' + ' '.join(texts))
        return _truncate_reference_text('\n'.join(slides))
    except Exception:
        return ''


def _extract_xlsx_text(path: Path) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True, data_only=True)
        chunks = []
        for ws in wb.worksheets[:5]:
            chunks.append(f'## 工作表：{ws.title}')
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                vals = [str(v).strip() for v in row if v not in (None, '')]
                if not vals:
                    continue
                chunks.append(' | '.join(vals[:12]))
                row_count += 1
                if row_count >= 120:
                    break
        return _truncate_reference_text('\n'.join(chunks))
    except Exception:
        return ''


def _extract_pdf_text(path: Path) -> str:
    cmds = [
        ['pdftotext', '-layout', str(path), '-'],
        ['strings', '-n', '6', str(path)],
    ]
    for cmd in cmds:
        try:
            out = subprocess.check_output(cmd, text=True, stderr=subprocess.DEVNULL)
            out = _truncate_reference_text(out)
            if len(out.strip()) >= 80:
                return out
        except Exception:
            pass
    return ''


def _extract_reference_file_text(path: str) -> str:
    pp = Path(path)
    if not pp.exists() or not pp.is_file():
        return ''

    ext = pp.suffix.lower()
    if ext in ('.txt', '.md', '.markdown', '.csv', '.tsv', '.rtf'):
        return _read_plain_text_file(pp)
    if ext == '.docx':
        return _extract_docx_text(pp)
    if ext == '.pptx':
        return _extract_pptx_text(pp)
    if ext in ('.xlsx', '.xlsm'):
        return _extract_xlsx_text(pp)
    if ext == '.pdf':
        return _extract_pdf_text(pp)
    return ''


def _ingest_reference_urls(metadata: dict, reference_urls: list) -> list:
    ingested = []
    ref_dir = CONFIG.get('reference_dir') or (CONFIG['output_dir'] / 'references')
    ref_dir.mkdir(parents=True, exist_ok=True)

    for url in reference_urls or []:
        try:
            if _is_notion_url(url):
                text, meta = _fetch_notion_reference_text(url)
                title = meta.get('title') or 'Notion 筆記'
                if _append_reference_section(metadata, f'Notion 筆記：{title}', text):
                    ingested.append(f'Notion:{title}')

                for section in meta.get('transcription_sections') or []:
                    summary = (section.get('summary') or '').strip()
                    transcript = (section.get('transcript') or '').strip()
                    content = (section.get('content') or '').strip()
                    if summary and _append_reference_section(metadata, f'Notion AI 摘要：{title}', summary):
                        ingested.append(f'NotionAI摘要:{title}')
                    if transcript:
                        if metadata.get('skip_notion_ai_transcript'):
                            if _store_notion_ai_transcript_candidate(metadata, title, transcript):
                                ingested.append(f'Notion逐字稿(保留備援):{title}')
                        elif _append_reference_section(metadata, f'Notion AI 逐字稿：{title}', transcript):
                            ingested.append(f'Notion逐字稿:{title}')
                    elif content and '逐字稿' in content:
                        if metadata.get('skip_notion_ai_transcript'):
                            _store_notion_ai_transcript_candidate(metadata, title, content)
                        else:
                            _append_reference_section(metadata, f'Notion 轉錄內容：{title}', content)

                media_blocks = meta.get('media_blocks') or []
                image_items = []
                file_items = []
                for item in media_blocks:
                    if not isinstance(item, dict):
                        continue
                    if _classify_notion_media_item(item) == 'image':
                        image_items.append(item)
                    elif _classify_notion_media_item(item) == 'document':
                        file_items.append(item)

                image_paths = []
                for idx, item in enumerate(image_items, 1):
                    media_url = (item.get('url') or '').strip()
                    if not media_url:
                        continue
                    try:
                        p = _download_notion_media(media_url, ref_dir, hint=f'{meta.get("page_id") or "page"}_{idx}')
                        image_paths.append(str(p))
                    except Exception as e:
                        logger.warning(f'下載 Notion 圖片失敗 {media_url}: {e}')

                ocr_rows = _ocr_image_texts(image_paths)
                _record_ocr_stats(metadata, f'Notion 圖片 OCR：{title}', image_paths, ocr_rows)
                ocr_text = _maybe_ocr_reference_images(image_paths, title)
                if ocr_text:
                    _append_reference_section(metadata, f'Notion 圖片 OCR：{title}', ocr_text)

                for idx, item in enumerate(file_items, 1):
                    media_url = (item.get('url') or '').strip()
                    if not media_url:
                        continue
                    hint = f'{meta.get("page_id") or "page"}_{idx}'
                    try:
                        p = _download_notion_media(media_url, ref_dir, hint=hint)
                        text = _extract_reference_file_text(str(p))
                        fname = p.name
                        if text:
                            _append_reference_section(metadata, f'Notion 附件：{fname}', text)
                            ingested.append(f'Notion附件:{fname}')
                        else:
                            _append_reference_section(metadata, f'Notion 附件：{fname}', f'已收到檔案 {fname}，但目前無法自動抽取文字；摘要時請至少將此檔視為背景參考資料。')
                            ingested.append(f'Notion附件:{fname}')
                    except Exception as e:
                        logger.warning(f'下載/抽取 Notion 附件失敗 {media_url}: {e}')
                        fname = (item.get('filename_hint') or Path(media_url.split("?", 1)[0]).name or 'notion_attachment')
                        _append_reference_section(metadata, f'Notion 附件：{fname}', f'已收到檔案 {fname}，但下載或抽取失敗；摘要時請至少將此檔視為背景參考資料。')
                continue

            if _looks_like_reference_doc_url(url):
                p = _download_reference_file(url, ref_dir)
                text = _extract_reference_file_text(str(p))
                title = p.name
                if text:
                    if _append_reference_section(metadata, f'教材附件：{title}', text):
                        ingested.append(f'附件:{title}')
                else:
                    if _append_reference_section(metadata, f'教材附件：{title}', f'已收到檔案 {title}，但目前無法自動抽取文字；摘要時請至少將此檔視為背景參考資料。'):
                        ingested.append(f'附件:{title}')
        except Exception as e:
            logger.warning(f'載入參考 URL 失敗 {url}: {e}')
    return ingested


def _ingest_reference_files(metadata: dict, file_paths: list) -> list:
    ingested = []
    for fp in file_paths or []:
        try:
            p = Path(fp)
            text = _extract_reference_file_text(str(p))
            title = p.name
            if text:
                if _append_reference_section(metadata, f'教材附件：{title}', text):
                    ingested.append(f'附件:{title}')
            else:
                if _append_reference_section(metadata, f'教材附件：{title}', f'已收到檔案 {title}，但目前無法自動抽取文字；摘要時請至少將此檔視為背景參考資料。'):
                    ingested.append(f'附件:{title}')
        except Exception as e:
            logger.warning(f'載入參考檔失敗 {fp}: {e}')
    return ingested


def _extract_reference_notes(text: str) -> str:
    """Extract user notes / reference bullets embedded in message text.

    Supports patterns like:
    - 以下是我的筆記 ...
    - 我的筆記:
    - 可作為參考 / 作為摘要參考
    """
    t = (text or '').strip()
    if not t:
        return ''

    markers = [
        '以下是我的筆記', '以下是 第一場演講 我個人的筆記', '以下是第二場演講 我個人的筆記',
        '以下是第三場演講 我個人的筆記', '我的筆記', '個人筆記', '筆記如下', '可作為參考', '作為摘要製作參考',
        '特記事項', '特殊要求', '摘要要求', '補充要求'
    ]
    pos = None
    for m in markers:
        idx = t.find(m)
        if idx != -1 and (pos is None or idx < pos):
            pos = idx
    if pos is None:
        mm = re.search(r'(^|\n)#{1,6}\s*.*筆記.*\n', t)
        if mm:
            pos = mm.start()
    if pos is None:
        return ''

    notes = t[pos:].strip()
    notes = re.sub(r'^(以下是[^\n]*筆記[^\n]*[:：]?\s*)', '', notes)
    notes = re.sub(r'^(我的筆記|個人筆記|筆記如下|特記事項|特殊要求|摘要要求|補充要求)\s*[:：]?\s*', '', notes)
    return notes.strip()


def _extract_user_fields(text: str) -> dict:
    """Best-effort parse for user-provided context in caption/message text."""
    if not (text or '').strip():
        return {}
    t = text.strip()
    out = {}
    # Patterns like: 主題: xxx
    for key in ["摘要主題", "主題", "課程", "課名", "課程名稱", "教授", "老師", "日期", "時間", "日期/時間", "地點", "人員", "與會", "參與", "參加", "DB", "資料庫", "模板", "Template", "template", "flag", "flags", "Flag", "Flags", "skip_notion_ai_transcript"]:
        mm = re.search(rf"{re.escape(key)}\s*[:：]\s*(.+)", t)
        if mm:
            out[key] = mm.group(1).strip()
    return out


def _truthy_flag_value(value: str) -> bool:
    v = (value or '').strip().lower()
    return v in {'1', 'true', 'yes', 'y', 'on', 'enable', 'enabled', 'skip_notion_ai_transcript'}


def _has_runtime_flag(provided: dict, ctx_text: str, flag_name: str) -> bool:
    provided = provided or {}
    flag_l = (flag_name or '').strip().lower()
    if not flag_l:
        return False

    direct = provided.get(flag_name)
    if isinstance(direct, str) and _truthy_flag_value(direct):
        return True

    for key in ('flag', 'flags', 'Flag', 'Flags'):
        raw = (provided.get(key) or '').strip().lower()
        if not raw:
            continue
        tokens = [tok.strip() for tok in re.split(r'[\s,，、/｜|;；]+', raw) if tok.strip()]
        if flag_l in tokens:
            return True

    return bool(re.search(rf'(?<![\w-]){re.escape(flag_name)}(?![\w-])', ctx_text or '', flags=re.I))


def _apply_runtime_flags(metadata: dict, provided: dict, ctx_text: str) -> dict:
    if _has_runtime_flag(provided, ctx_text, 'skip_notion_ai_transcript'):
        metadata['skip_notion_ai_transcript'] = True
    return metadata


def _parse_date_time_answer(ans: str, fallback_date: str) -> Tuple[str, str]:
    """Return (date, time). Accepts: YYYY-MM-DD HH:MM | YYYY/MM/DD HH:MM | HH:MM | YYYY-MM-DD."""
    a = (ans or '').strip()
    if not a:
        return fallback_date, ""

    a = a.replace('/', '-')
    # full datetime
    m = re.match(r"^(\d{4}-\d{2}-\d{2})\s+(\d{1,2}:\d{2})$", a)
    if m:
        return m.group(1), m.group(2)
    # date only
    m = re.match(r"^(\d{4}-\d{2}-\d{2})$", a)
    if m:
        return m.group(1), ""
    # time only
    m = re.match(r"^(\d{1,2}:\d{2})$", a)
    if m:
        return fallback_date, m.group(1)

    return fallback_date, a


def _suggest_db_type(recording_time, topic: str) -> str:
    """Return suggested type: 'emba' or 'business'."""
    t = (topic or '').lower()
    if any(k in t for k in ["emba", "課堂", "教授", "個案", "作業", "期末"]):
        return 'emba'
    if recording_time is not None:
        inf = _infer_course(recording_time)
        if inf.get('auto_detected'):
            return 'emba'
    return 'business'


async def _ask_common_context(send_message, ask_user, metadata: dict, provided: dict) -> dict:
    """Ask for topic/date-time/location/people when missing."""
    # 1) topic
    topic = provided.get('摘要主題') or provided.get('主題')
    if not (topic or '').strip():
        await send_message("我需要幾個資訊幫尼把摘要整理得更準：\n1) 摘要主題  2) 日期/時間  3) 地點  4) 人員")
        topic = await ask_user("摘要主題")
    topic = (topic or '').strip()

    # 2) date/time
    dt_ans = provided.get('日期/時間') or ''
    if not dt_ans:
        # allow separate date/time
        d = (provided.get('日期') or '').strip()
        tm = (provided.get('時間') or '').strip()
        if d and tm:
            dt_ans = f"{d} {tm}"
        elif d:
            dt_ans = d
        elif tm:
            dt_ans = tm

    if not (dt_ans or '').strip():
        dt_ans = await ask_user(f"日期/時間（可只回時間；預設 {metadata.get('date','')}）")

    date, tm = _parse_date_time_answer(dt_ans, metadata.get('date',''))
    metadata['date'] = date
    if tm:
        metadata['time'] = tm

    # 3) location
    loc = provided.get('地點') or metadata.get('location') or metadata.get('room')
    if not (loc or '').strip():
        loc = await ask_user("地點")
    metadata['location'] = (loc or '').strip()

    # 4) people
    ppl = provided.get('人員') or provided.get('與會') or provided.get('參與') or ''
    if not (ppl or '').strip():
        ppl = await ask_user("人員（逗號分隔；不知道可回 ok）")
    ppl = (ppl or '').strip()
    people_list = []
    if ppl.lower() not in ('ok','跳過',''):
        people_list = [x.strip() for x in ppl.split(',') if x.strip()]

    metadata['attendees'] = people_list
    metadata['participants'] = people_list

    # store topic in both keys (downstream compatibility)
    metadata.setdefault('course_name', '')
    metadata.setdefault('meeting_name', '')
    if not metadata.get('course_name'):
        metadata['course_name'] = topic
    if not metadata.get('meeting_name'):
        metadata['meeting_name'] = topic

    return metadata

# ============================================================
# Markdown 生成
# ============================================================

def _build_transcript_md(segments, speakers):
    """逐字稿 Markdown (toggle heading 用)"""
    main_spk = next((s for s, i in speakers.items() if i.get("is_main_speaker")), None)
    lines = []

    if speakers:
        total = sum(s["duration_sec"] for s in speakers.values())
        lines.extend(["## 說話者統計", "",
                       "| 說話者 | 角色 | 發言時長 | 比例 |",
                       "|--------|------|----------|------|"])
        for spk, info in sorted(speakers.items(), key=lambda x: x[1]["duration_sec"], reverse=True):
            name = info.get("display_name", spk)
            pct = f"{info['duration_sec']/total*100:.0f}%" if total else "?"
            lines.append(f"| {name} | {info.get('role','')} | {_fmt_dur(info['duration_sec'])} | {pct} |")

    lines.extend(["", "## 完整逐字稿", ""])
    cur_spk = None
    for seg in segments:
        spk = seg.get("speaker", "")
        text = seg.get("text", "").strip()
        if not text:
            continue
        if spk and spk != cur_spk:
            cur_spk = spk
            name = speakers.get(spk, {}).get("display_name", spk)
            lines.extend(["", f"**{name}** `[{_fmt_ts(seg.get('start', 0))}]`", ""])
        lines.append(text)
    return "\n".join(lines)


def _build_transcript_plain(segments, speakers):
    """純文字逐字稿 (LLM 摘要用)"""
    lines = []
    for seg in segments:
        spk = seg.get("speaker", "")
        name = speakers.get(spk, {}).get("display_name", spk) if spk else ""
        text = seg.get("text", "").strip()
        ts = _fmt_ts(seg.get("start", 0))
        lines.append(f"[{ts}] {name}: {text}" if name else f"[{ts}] {text}")
    return "\n".join(lines)


def _chunk_segments_for_llm(segments, speakers, max_chars: int = 12000, max_minutes: int = 12):
    """Split transcript into chunk texts for LLM.

    Design goals:
    - keep each chunk small enough for fast/cheap models
    - avoid LLM context explosion for multi-hour audio
    """
    chunks = []
    buf = []
    chunk_start = None
    last_ts = 0

    def flush():
        nonlocal buf, chunk_start, last_ts
        if not buf:
            return
        start_ts = _fmt_ts(chunk_start or 0)
        end_ts = _fmt_ts(last_ts or (chunk_start or 0))
        label = f"{start_ts}-{end_ts}"
        chunks.append({"label": label, "text": "\n".join(buf)})
        buf = []
        chunk_start = None

    for seg in segments:
        text = (seg.get("text") or "").strip()
        if not text:
            continue
        st = int(seg.get("start") or 0)
        last_ts = st
        if chunk_start is None:
            chunk_start = st

        spk = seg.get("speaker", "")
        name = speakers.get(spk, {}).get("display_name", spk) if spk else ""
        ts = _fmt_ts(st)
        line = f"[{ts}] {name}: {text}" if name else f"[{ts}] {text}"

        # flush by time window
        if chunk_start is not None and (st - chunk_start) >= max_minutes * 60 and buf:
            flush()
            chunk_start = st

        # flush by char budget
        if sum(len(x) for x in buf) + len(line) + 1 > max_chars and buf:
            flush()
            chunk_start = st

        buf.append(line)

    flush()
    return chunks



# ============================================================
# 類別正規化
# ============================================================

def _norm_category(t):
    m = {"出差": "出差🟠", "來訪": "來訪🔵", "線上": "線上🟣",
         "online": "線上🟣", "visit": "來訪🔵"}
    return m.get(t.strip().lower(), t.strip())

def _norm_dept(t):
    m = {"ws": "WS業務🟢", "wq": "WQ品管🔴", "跨部門": "跨部門🟡",
         "業務": "WS業務🟢", "品管": "WQ品管🔴"}
    return m.get(t.strip().lower(), t.strip())

def _extract_action_items(md):
    items = []
    for line in md.split("\n"):
        s = line.strip()
        if s.startswith("- [ ]") or s.startswith("- [x]"):
            items.append(s[6:].strip())
    return "\n".join(f"{i+1}. {it}" for i, it in enumerate(items)) if items else ""



def _extract_report_lines(summary: str):
    """Extract numbered report lines for business report code block.

    Priority:
    1) Lines under '### 業務會報版摘要' section (if present)
    2) Any numbered list lines '1. ...' in the whole summary (for template R)
    """
    if not summary:
        return []

    lines = summary.split('\n')

    # 1) Prefer section-based extraction
    out = []
    in_sec = False
    for ln in lines:
        s = ln.strip()
        if s.startswith('### ') and '業務會報版摘要' in s:
            in_sec = True
            continue
        if in_sec and s.startswith('### '):
            break
        if in_sec and re.match(r'^\d+\.', s):
            out.append(s)
    if out:
        return out

    # 2) Fallback: grab any numbered list lines
    for ln in lines:
        s = ln.strip()
        if re.match(r'^\d+\.', s):
            out.append(s)
    return out
def _fmt_dur(sec):
    m, s = divmod(int(sec), 60)
    h, m = divmod(m, 60)
    return f"{h}h {m:02d}m {s:02d}s" if h else f"{m}m {s:02d}s"

def _fmt_ts(sec):
    m, s = divmod(int(sec), 60)
    h, m = divmod(m, 60)
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


# ============================================================
# LLM 呼叫
# ============================================================

def _gemini_direct_model_name(model_id: str) -> str:
    mm = (model_id or '').strip()
    if mm in _GEMINI_DIRECT_MODEL_MAP:
        return _GEMINI_DIRECT_MODEL_MAP[mm]
    if mm.startswith('google/'):
        return mm.split('/', 1)[1]
    return ''


def _gemini_direct_completion(model_id: str, system_prompt: str, user_message: str, max_tokens: int = 8000) -> str:
    api_key = (os.environ.get('GEMINI_API_KEY') or os.environ.get('GOOGLE_API_KEY') or '').strip()
    if not api_key:
        raise RuntimeError('GEMINI_API_KEY / GOOGLE_API_KEY 未設定，無法使用 Gemini direct fallback')

    model_name = _gemini_direct_model_name(model_id)
    if not model_name:
        raise RuntimeError(f'模型 {model_id} 無 Gemini direct fallback')

    url = f'https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}'
    payload = {
        'system_instruction': {'parts': [{'text': system_prompt}]},
        'contents': [{'parts': [{'text': user_message}]}],
        'generationConfig': {
            'temperature': 0.2,
            'topP': 0.95,
            'maxOutputTokens': max_tokens,
        },
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode('utf-8'),
        headers={'Content-Type': 'application/json'},
    )
    with urllib.request.urlopen(req, timeout=300) as resp:
        data = json.loads(resp.read().decode('utf-8'))
    parts = (((data.get('candidates') or [{}])[0].get('content') or {}).get('parts') or [])
    texts = [p.get('text', '') for p in parts if isinstance(p, dict) and p.get('text') and not p.get('thought')]
    text = '\n'.join(texts).strip()
    if not text:
        raise RuntimeError(f'Gemini direct completion returned empty content for {model_name}')
    return text


async def _call_llm(
    model_id,
    system_prompt,
    user_message,
    send_message=None,
    max_tokens: int = 8000,
    retries: int = 3,
    fallback_chain=None,
    usage_bucket: str = 'final',
    proc: Optional[dict] = None,
):
    """
    呼叫 LLM。

    - final / chunk 可各自帶不同 fallback chain
    - fallback 啟動時主動通知使用者
    - 記錄實際成功使用的模型，供 footer / Notion / Obsidian 寫入
    """
    import asyncio

    def _dedupe_models(models):
        out = []
        seen = set()
        for m in models:
            mm = (m or '').strip()
            if not mm or mm in seen:
                continue
            seen.add(mm)
            out.append(mm)
        return out

    def _is_fallbackworthy_error(err: Exception) -> bool:
        msg = str(err or '').lower()
        keys = [
            'credit balance is too low', 'insufficient', 'quota', 'rate limit', '429',
            'overloaded', 'capacity', 'temporarily unavailable', 'timeout', 'timed out',
            'connection', 'api key', 'authentication', 'unauthorized', 'forbidden',
            '529',
        ]
        return any(k in msg for k in keys)

    def _append_unique(lst, value):
        if value and value not in lst:
            lst.append(value)

    def _record_success(requested_model: str, used_model: str):
        if not isinstance(proc, dict):
            return
        usage = proc.setdefault('llm_usage', {})
        bucket = usage.setdefault(usage_bucket, {})
        bucket['requested_model'] = requested_model
        bucket['final_used_model'] = used_model
        _append_unique(bucket.setdefault('success_models', []), used_model)
        bucket['used_fallback'] = (used_model != requested_model)
        if fallback_chain:
            bucket['fallback_chain'] = list(fallback_chain)
        if used_model.startswith('google/') and used_model != requested_model:
            _append_proc_warning(proc, f'⚠️ {usage_bucket} 使用 Google 付費 fallback：{used_model}')
        if used_model.startswith('google/') and bucket.get('direct_fallback') == 'gemini-api':
            _append_proc_warning(proc, f'⚠️ {usage_bucket} 使用 Gemini direct API fallback：{used_model}')

    async def _notify(msg: str):
        if not send_message:
            return
        try:
            await send_message(msg)
        except Exception:
            pass

    async def _notify_once(key: str, msg: str):
        if isinstance(proc, dict):
            sent = proc.setdefault('_fallback_notice_keys', [])
            if key in sent:
                return
            sent.append(key)
        await _notify(msg)

    model_chain = _dedupe_models([model_id] + list(fallback_chain or []))

    # 方式 1: OpenClaw routing（首選，支援 OAuth / provider routing）
    try:
        from openclaw.llm import chat_completion
        last_err = None
        for idx, candidate in enumerate(model_chain):
            for attempt in range(retries):
                try:
                    if idx > 0 and attempt == 0:
                        await _notify_once(
                            f'{usage_bucket}:fallback-start',
                            f"⚠️ {usage_bucket} 主模型 {model_chain[0]} 暫時不可用，已啟動 fallback 機制。",
                        )
                        note = f"↪️ lecture-transcribe {usage_bucket} 模型 fallback → {candidate}"
                        logger.warning(note)
                        await _notify_once(f'{usage_bucket}:candidate:{candidate}', note)
                    resp = await chat_completion(
                        model=candidate,
                        system=system_prompt,
                        messages=[{"role": "user", "content": user_message}],
                        max_tokens=max_tokens,
                    )
                    _record_success(model_id, candidate)
                    return resp.get("content", "")
                except Exception as e:
                    last_err = e
                    should_retry_same = attempt < retries - 1 and _is_fallbackworthy_error(e)
                    if should_retry_same:
                        await asyncio.sleep(min(20, 4 ** attempt))
                        continue
                    break
            # Delay between fallback models to let overloaded API recover
            if idx > 0:
                await asyncio.sleep(min(30, 10 * idx))
        raise last_err
    except ImportError as e:
        if not ALLOW_GOOGLE_DIRECT_FALLBACK:
            chain_text = " → ".join(model_chain) if model_chain else (model_id or "(未指定)")
            msg = (
                f"OpenClaw routing 不可用，且 lecture-transcribe 已停用 Gemini direct fallback。"
                f" 可用模型鏈：{chain_text}"
            )
            _append_proc_warning(proc, f'⚠️ {usage_bucket} routing miss，Gemini direct fallback 預設停用')
            await _notify_once(
                f'{usage_bucket}:routing-miss-blocked',
                f"⚠️ OpenClaw routing 不可用，lecture-transcribe {usage_bucket} 不會改走 Gemini direct fallback。",
            )
            raise RuntimeError(msg) from e
        await _notify_once(
            f'{usage_bucket}:routing-miss',
            f"⚠️ OpenClaw routing 不可用，lecture-transcribe {usage_bucket} 已明確允許 Gemini direct fallback，將只在 Google 模型上嘗試。",
        )
        chain_text = " → ".join(model_chain) if model_chain else (model_id or "(未指定)")
        logger.warning(f"OpenClaw routing 不可用，改用 direct fallback 嘗試模型鏈：{chain_text}")
        _append_proc_warning(proc, f'⚠️ {usage_bucket} routing miss，已啟用 Gemini direct fallback')
        last_err = e
        for idx, candidate in enumerate(model_chain):
            if not (candidate or '').startswith('google/'):
                last_err = RuntimeError(f'模型 {candidate} 無 direct fallback，已略過')
                continue
            for attempt in range(retries):
                try:
                    if (idx > 0 or candidate != model_id) and attempt == 0:
                        await _notify_once(
                            f'{usage_bucket}:direct-candidate:{candidate}',
                            f"↪️ lecture-transcribe {usage_bucket} direct fallback → {candidate}",
                        )
                    import asyncio
                    resp_text = await asyncio.to_thread(
                        _gemini_direct_completion,
                        candidate,
                        system_prompt,
                        user_message,
                        max_tokens,
                    )
                    _record_success(model_id, candidate)
                    usage = proc.setdefault('llm_usage', {}) if isinstance(proc, dict) else {}
                    bucket = usage.setdefault(usage_bucket, {}) if isinstance(usage, dict) else {}
                    if isinstance(bucket, dict):
                        bucket['direct_fallback'] = 'gemini-api'
                    _append_proc_warning(proc, f'⚠️ {usage_bucket} 使用 Gemini direct API fallback：{candidate}')
                    return resp_text
                except Exception as inner:
                    last_err = inner
                    should_retry_same = attempt < retries - 1 and _is_fallbackworthy_error(inner)
                    if should_retry_same:
                        await asyncio.sleep(min(20, 4 ** attempt))
                        continue
                    break
        raise RuntimeError(
            f"OpenClaw routing 不可用，且 direct fallback 也失敗。可用模型鏈：{chain_text} | 最後錯誤：{last_err}"
        ) from e


# ============================================================
# 主流程
# ============================================================

async def handle_audio_message(
    audio_path, send_message, ask_user,
    recording_time=None, pending_files=None,
    message_text: Optional[str] = None,
    caption: Optional[str] = None,
    **kwargs,
):
    """
    完整流程入口。由 OpenClaw Telegram handler 呼叫。
    """
    def progress(msg):
        import asyncio
        try:
            loop = asyncio.get_running_loop()
            loop.create_task(send_message(msg))
        except RuntimeError:
            logger.info(msg)

    # ===== 多檔合併 =====
    if pending_files and len(pending_files) > 1:
        file_list = "\n".join(
            f"  {i+1}. {Path(f).name} ({get_audio_duration(f)/60:.0f}min)"
            for i, f in enumerate(pending_files)
        )
        await send_message(
            f"🎙️ 偵測到 {len(pending_files)} 個音訊檔:\n{file_list}\n\n"
            f"  1️⃣ 合併為同一場\n  2️⃣ 分開處理"
        )
        r = await ask_user("合併?")
        if r.strip() in ("1", "合併"):
            merged = str(CONFIG["output_dir"] / "merged_audio.m4a")
            CONFIG["output_dir"].mkdir(parents=True, exist_ok=True)
            audio_path = merge_audio_files(pending_files, merged)
            await send_message(f"✅ 合併完成 ({get_audio_duration(audio_path)/60:.0f}min)")
        else:
            audio_path = pending_files[0]
            await send_message("👌 先處理第一段")

    # ===== Step 1: 類型 + Metadata =====
    dur_sec = get_audio_duration(audio_path)
    dur_min = dur_sec / 60

    # Base metadata (date/day-of-week) inferred from recording_time or now
    metadata = _infer_course(recording_time)

    # Parse any user-provided context from caption/message
    ctx_text = (caption or '') + "\n" + (message_text or '')
    provided = _extract_user_fields(ctx_text)
    metadata = _apply_user_context_overrides(metadata, provided, ctx_text)
    metadata = _apply_runtime_flags(metadata, provided, ctx_text)

    reference_notes = _extract_reference_notes(ctx_text)
    if reference_notes:
        _append_reference_section(metadata, '使用者補充筆記', reference_notes)

    reference_urls = kwargs.get('reference_urls') or _pick_reference_urls(ctx_text)
    reference_file_paths = kwargs.get('reference_file_paths') or []
    reference_file_paths = [p for p in reference_file_paths if str(p) != str(audio_path)]

    ingested_sources = []
    if reference_urls:
        ingested_sources.extend(_ingest_reference_urls(metadata, reference_urls[:6]))
    if reference_file_paths:
        ingested_sources.extend(_ingest_reference_files(metadata, reference_file_paths[:8]))

    image_paths = kwargs.get('image_paths') or []
    image_ocr = _ocr_image_texts(image_paths)
    _record_ocr_stats(metadata, '使用者提供圖片 OCR', image_paths, image_ocr)
    if image_ocr:
        ocr_text = _maybe_ocr_reference_images(image_paths, '使用者提供教材照片')
        if ocr_text:
            metadata['image_ocr_notes'] = _truncate_ocr_reference_text(ocr_text, max_chars=6000)
            if _append_reference_section(metadata, '教材照片 OCR：使用者提供圖片', metadata['image_ocr_notes']):
                ingested_sources.append(f'圖片:{len(image_paths)}')

    reference_entities = _build_reference_entities(metadata, ctx_text + "\n" + (metadata.get('image_ocr_notes') or ''))
    if reference_entities:
        metadata['reference_entities'] = reference_entities

    if ingested_sources:
        pretty = '、'.join(ingested_sources[:6])
        more = '' if len(ingested_sources) <= 6 else f' 等 {len(ingested_sources)} 項'
        await send_message(f"📚 已整合參考資料：{pretty}{more}")

    # Allow explicit override: even if course schedule matched, user may want business DB (e.g., forum talks used for EMBA homework)
    if metadata.get("auto_detected"):
        t_hint = (provided.get('DB') or provided.get('資料庫') or '').lower()
        template_hint = (provided.get('模板') or provided.get('Template') or provided.get('template') or '').strip().upper()
        if any(k in t_hint for k in ("商務", "business", "會談", "meeting")) or template_hint in ("E",):
            metadata["auto_detected"] = False
            metadata["auto_detected_override"] = "user_forced_business"

    enough_emba_context = bool(metadata.get('course_name') and metadata.get('professor'))

    if metadata["auto_detected"] or metadata.get('schedule_matched_by_user_context'):
        # EMBA 自動（課表命中）
        # 仍允許補齊地點/人員/時間等資訊
        if not enough_emba_context:
            await send_message(
                f"📚 EMBA 課堂:\n"
                f"  {metadata['course_name']} / {metadata['professor']}\n"
                f"  {metadata['date']} 週{metadata['day_of_week']} / {dur_min:.0f}min\n\n"
                f"我再跟尼確認一下：地點/人員/時間有需要補嗎？（沒有就回 ok）"
            )
            extra = await ask_user("補充")
            if (extra or '').lower() not in ("ok", "跳過", ""):
                if "," in extra:
                    metadata["attendees"] = [x.strip() for x in extra.split(",") if x.strip()]
                    metadata["participants"] = metadata["attendees"]
                else:
                    metadata["location"] = extra.strip()

            await send_message("補充關鍵字? (逗號分隔 或 ok)")
            kw = await ask_user("關鍵字")
            metadata["keywords"] = (
                [k.strip() for k in (kw or '').split(",") if k.strip()]
                if (kw or '').lower() not in ("ok", "跳過", "") else []
            )
        else:
            metadata.setdefault('keywords', [])
        metadata["model_pref"] = None
        metadata["type"] = "emba"

    else:
        # 使用者沒主動給錄音資訊時：主動問『摘要需要的 4 個欄位』
        # 並在需要 DB 路由（課堂 vs 商務）時，先詢問要上傳到哪個 DB。
        await send_message(
            f"📝 收到錄音 ({metadata['date']} 週{metadata['day_of_week']}, {dur_min:.0f}min)"
        )

        # 5) DB 選擇（以錄音上傳時間 + 主題做建議，但仍請使用者確認）
        chosen = None
        t_hint = (provided.get('DB') or provided.get('資料庫') or '').lower()
        if any(k in t_hint for k in ("課堂", "emba", "class")):
            chosen = 'emba'
        elif any(k in t_hint for k in ("商務", "會談", "meeting", "business")):
            chosen = 'business'

        if not chosen:
            # Need topic first for a better suggestion
            if not (provided.get('摘要主題') or provided.get('主題')):
                await send_message("我先問一句：這段錄音的『摘要主題』是？")
                provided['主題'] = (await ask_user("摘要主題")).strip()

            suggest = _suggest_db_type(recording_time, provided.get('摘要主題') or provided.get('主題') or '')
            sug_label = '課堂摘要DB' if suggest == 'emba' else '商務會談DB'
            await send_message(
                "📤 這份摘要要上傳到哪個 Notion DB？\n"
                "  1️⃣ 📚 課堂摘要DB\n"
                "  2️⃣ 💼 商務會談DB\n\n"
                f"我猜是：{sug_label}（但我想跟尼確認）"
            )
            db = await ask_user("上傳DB")
            if db.strip() in ("1", "課堂", "emba"):
                chosen = 'emba'
            elif db.strip() in ("2", "商務", "會談", "business", ""):
                chosen = 'business'
            else:
                chosen = 'other'

        metadata['type'] = chosen

        # 1-4) 補齊摘要所需資訊
        metadata = await _ask_common_context(send_message, ask_user, metadata, provided)

        # 類型專屬補充
        if metadata['type'] == 'emba':
            await send_message("再補兩個欄位就好：課程名稱 / 教授（不確定可回 ok）")
            info = await ask_user("課程/教授")
            parts = [p.strip() for p in (info or '').split("/")]
            if parts and parts[0] and parts[0].lower() not in ('ok','跳過'):
                metadata['course_name'] = parts[0]
                metadata['meeting_name'] = parts[0]
            if len(parts) > 1 and parts[1] and parts[1].lower() not in ('ok','跳過'):
                metadata['professor'] = parts[1]

            t_in = (provided.get('模板') or provided.get('Template') or provided.get('template') or '').strip()
            if t_in:
                metadata['template_override'] = t_in.strip().upper()
            # keywords optional
            await send_message("關鍵字(選填，逗號分隔；沒有就回 ok)")
            kw = await ask_user("關鍵字")
            metadata["keywords"] = (
                [k.strip() for k in (kw or '').split(",") if k.strip()]
                if (kw or '').lower() not in ("ok", "跳過", "") else []
            )
            metadata['model_pref'] = None

        elif metadata['type'] == 'business':
            # Template hint (talk/forum)
            topic_hint = (metadata.get('topic') or provided.get('摘要主題') or provided.get('主題') or '').lower()
            if any(k in topic_hint for k in ("演講", "論壇", "講座", "keynote", "speech", "panel")):
                metadata['template_override'] = 'E'

            # Allow user to force template via caption/message: 模板: E
            t_in = (provided.get('模板') or provided.get('Template') or provided.get('template') or '').strip()
            if t_in:
                metadata['template_override'] = t_in.strip().upper()

            await send_message(
                "（選填，但會讓 Notion 欄位更完整）回覆 / 分隔：\n"
                "  對象公司 / 類別(出差/來訪/線上) / 課別(WS/WQ/跨部門) / 模型(選填)\n\n"
                "例: SUBARU-JP / 來訪 / WS\n\n"
                "另外：如果這次是『論壇/演講活動』摘要，尼可以回我：模板 E（不回就自動判斷）"
            )
            biz2 = await ask_user("補充")
            p2 = [x.strip() for x in (biz2 or '').split("/")]
            if len(p2) > 0 and p2[0] and p2[0].lower() not in ('ok','跳過'):
                metadata['company'] = p2[0]
                metadata['target_company'] = p2[0]
            if len(p2) > 1 and p2[1] and p2[1].lower() not in ('ok','跳過'):
                metadata['category'] = _norm_category(p2[1])
            if len(p2) > 2 and p2[2] and p2[2].lower() not in ('ok','跳過'):
                metadata['department'] = _norm_dept(p2[2])
            if len(p2) > 3 and p2[3] and p2[3].lower() not in ('ok','跳過'):
                metadata['model_pref'] = p2[3].strip()

            # Ask template override (optional, default auto)
            if not metadata.get('template_override'):
                await send_message("模板(選填)：A/B/C/E（E=論壇/演講）。不指定就自動判斷，直接回 ok 也行。")
                tv = await ask_user("模板")
                if (tv or '').strip() and (tv or '').strip().lower() not in ('ok','跳過'):
                    metadata['template_override'] = (tv or '').strip().upper()

            metadata['keywords'] = []

        else:
            metadata['course_name'] = metadata.get('course_name') or '其他錄音'
            metadata['model_pref'] = None


    # ===== Step 2: 轉錄（含本地快取，避免重跑） =====
    await send_message("🎙️ 開始轉錄...")
    proc = metadata.setdefault('process_summary', {})
    proc['downloaded_from_cloud'] = bool(_pick_audio_url(ctx_text))
    proc['preprocess_tool'] = 'ffmpeg'
    proc['preprocess_desc'] = '16kHz 單聲道正規化與 loudnorm 音量校正'
    proc['reference_sources'] = metadata.get('reference_sources') or []
    proc['reference_source_count'] = len(proc['reference_sources'])

    cache_dir = CONFIG.get("cache_dir") or (CONFIG["output_dir"] / "cache")
    cache_dir.mkdir(parents=True, exist_ok=True)
    ap = Path(audio_path)
    # Use size+stem as cache key so re-downloaded files (same content) still hit cache
    cache_key = f"{ap.stem}__{ap.stat().st_size}.transcribe.json"
    cache_path = cache_dir / cache_key

    tx = None

    # For long audio (>45 min), skip whole-file cache and rely on per-chunk caching inside transcribe_audio
    if dur_sec > 45 * 60:
        proc['used_cache'] = False  # per-chunk caching handled internally
        t0 = time.time()
        await send_message(f"🎙️ 長音訊模式 ({dur_min:.0f}min)：啟用分段轉錄 + 每段快取...")
        try:
            tx = transcribe_audio(audio_path, progress_cb=progress, chunk_minutes=30)
        except Exception as e:
            await send_message(f"❌ 轉錄失敗: {e}")
            return
        elapsed = time.time() - t0
        engine = tx.get("engine_used", "")
        ratio = tx.get("duration_sec", 0) / elapsed if elapsed > 0 else 0
        proc['transcribe_engine'] = engine
        proc['segment_count'] = len(tx.get('segments', []) or [])
        proc['audio_duration_text'] = _fmt_dur(tx.get('duration_sec', 0))
        proc['transcribe_elapsed_sec'] = elapsed
        proc['transcribe_elapsed_text'] = f"{elapsed:.0f} 秒"
        await send_message(
            f"✅ 轉錄完成 | {engine} | {elapsed:.0f}s ({ratio:.1f}x)\n"
            f"  {_fmt_dur(tx.get('duration_sec', 0))} / {len(tx.get('segments', []) or [])} 段"
        )
    else:
        # Short audio: use existing whole-file cache logic
        if cache_path.exists():
            try:
                import json
                tx = json.loads(cache_path.read_text(encoding="utf-8"))
                if isinstance(tx, dict) and tx.get("segments"):
                    proc['used_cache'] = True
                    await send_message(f"♻️ 使用已轉錄快取：{cache_path.name}")
            except Exception:
                tx = None

        if tx is None:
            proc['used_cache'] = False
            t0 = time.time()
            try:
                tx = transcribe_audio(audio_path, progress_cb=progress)
            except Exception as e:
                await send_message(f"❌ 轉錄失敗: {e}")
                return

            # write cache for stability
            try:
                import json
                cache_path.write_text(json.dumps(tx, ensure_ascii=False), encoding="utf-8")
            except Exception:
                pass

            elapsed = time.time() - t0
            engine = tx.get("engine_used", "")
            ratio = tx["duration_sec"] / elapsed if elapsed > 0 else 0
            proc['transcribe_engine'] = engine
            proc['segment_count'] = len(tx.get('segments', []) or [])
            proc['audio_duration_text'] = _fmt_dur(tx.get('duration_sec', 0))
            proc['transcribe_elapsed_sec'] = elapsed
            proc['transcribe_elapsed_text'] = f"{elapsed:.0f} 秒"
            await send_message(
                f"✅ 轉錄完成 | {engine} | {elapsed:.0f}s ({ratio:.1f}x)\n"
                f"  {_fmt_dur(tx['duration_sec'])} / {len(tx['segments'])} 段"
            )
        else:
            engine = tx.get("engine_used", "(cached)")
            proc.setdefault('transcribe_engine', engine)
            proc.setdefault('segment_count', len(tx.get('segments', []) or []))
            proc.setdefault('audio_duration_text', _fmt_dur(tx.get('duration_sec', 0)))
            await send_message(
                f"✅ 轉錄完成 | {engine}\n"
                f"  {_fmt_dur(tx.get('duration_sec', 0))} / {len(tx.get('segments', []) or [])} 段"
            )

    # ===== Step 3: Diarization（預設關閉；需要才開） =====
    want_spk = _want_diarization(ctx_text)
    if want_spk:
        proc['used_diarization'] = True
        segs, speakers = run_diarization(audio_path, tx["segments"], progress)
        tx["segments"] = segs
    else:
        proc['used_diarization'] = False
        segs = tx["segments"]
        speakers = {}
        await send_message("🧍 Speaker 分辨（diarization）預設不執行：本次依照『準確度優先＋穩定』設定先略過。若尼想要分講者，下次訊息加上：speaker / 說話者 / 分辨 speaker。")

    # Speaker 預填 (階段 A)
    if speakers and metadata.get("attendees"):
        sorted_s = sorted(speakers, key=lambda s: speakers[s]["duration_sec"], reverse=True)
        for i, spk in enumerate(sorted_s):
            speakers[spk]["display_name"] = (
                metadata["attendees"][i] if i < len(metadata["attendees"]) else spk
            )

    # Speaker 校正 (階段 B)
    if speakers and len(speakers) > 1:
        # Build representative quotes per speaker (方案1)
        quotes_by = build_speaker_quotes(tx["segments"], speakers, top_n=3)
        kept, ignored = filter_non_topic_speakers(speakers, quotes_by)

        preview = get_speaker_preview({k: speakers[k] for k in kept}) if kept else get_speaker_preview(speakers)
        sorted_s = kept if kept else sorted(speakers, key=lambda s: speakers[s]["duration_sec"], reverse=True)

        msg = f"🎤 偵測到 {len(speakers)} 位說話者"
        if ignored:
            msg += f"（已先忽略 {len(ignored)} 位：談話量少/零散/多為閒聊）"
        msg += f":\n{preview}\n\n"

        # show quotes for kept speakers only
        for spk in sorted_s:
            q = (quotes_by.get(spk) or {}).get("quotes", [])
            if not q:
                continue
            disp = speakers[spk].get("display_name", spk)
            msg += f"【{spk} → {disp}】\n"
            for i, it in enumerate(q, 1):
                msg += f"  {i}. ({_fmt_ts(it['start'])}-{_fmt_ts(it['end'])}) {it['text']}\n"
            msg += "\n"

        if metadata.get("attendees"):
            msg += "目前對應:\n"
            for spk in sorted_s:
                msg += f"  {spk} → {speakers[spk].get('display_name', spk)}\n"

        msg += "\n修改: 00=名稱, 01=名稱\n或 ok 確認（若要把被忽略的也加回來，回覆：include 02,03…）"

        await send_message(msg)
        sr = await ask_user("Speaker")

        # Optional: allow user to include ignored speakers back
        sr_l = (sr or "").strip().lower()
        if sr_l.startswith("include") and ignored:
            # parse include indices like: include 02,03
            inc = re.findall(r"\b\d+\b", sr_l)
            for idx in inc:
                for spk in ignored:
                    if idx in spk and spk not in sorted_s:
                        sorted_s.append(spk)
            # re-sort with newly included speakers
            sorted_s = sorted(sorted_s, key=lambda s: speakers[s]["duration_sec"], reverse=True)

        if sr_l not in ("ok", "跳過", "") and not sr_l.startswith("include"):
            for pair in sr.split(","):
                if "=" in pair:
                    idx, name = pair.split("=", 1)
                    for spk in speakers:
                        if idx.strip() in spk:
                            speakers[spk]["display_name"] = name.strip()
                            break

    for spk in speakers:
        speakers[spk].setdefault("display_name", spk)

    # ===== Step 4: LLM 摘要 =====
    # Build transcript once (for local archive / glossary detection / template routing)
    plain = _build_transcript_plain(segs, speakers)
    trans_md = _build_transcript_md(segs, speakers)
    apply_course_specific_summary_flags(metadata, plain)
    transcript_quality = _assess_transcript_quality(segs, plain, tx.get('duration_sec', 0))
    proc['transcript_quality'] = transcript_quality
    proc['skip_notion_ai_transcript'] = bool(metadata.get('skip_notion_ai_transcript'))
    if metadata.get('skip_notion_ai_transcript'):
        proc['notion_ai_transcript_candidate_count'] = len(metadata.get('notion_ai_transcript_candidates') or [])
        proc['notion_ai_transcript_fallback_used'] = _maybe_enable_notion_ai_transcript_fallback(metadata, transcript_quality)
        proc['reference_sources'] = metadata.get('reference_sources') or proc.get('reference_sources') or []
        proc['reference_source_count'] = len(proc.get('reference_sources') or [])
    if metadata.get('enable_company_web_enrichment'):
        enrichment = _prepare_company_web_enrichment(metadata, plain)
        proc['company_candidate_count'] = len(metadata.get('company_candidates') or [])
        proc['company_enrichment_count'] = len(enrichment or [])

    spk_count = len(speakers) if speakers else 1
    tmpl = select_template(
        metadata.get("type", "business"),
        spk_count,
        tx["duration_sec"]/60,
        override=(metadata.get("template_override") or None),
        transcript_text=plain,
        metadata=metadata,
    )
    final_model = select_model(tx["duration_sec"]/60, metadata.get("model_pref"))

    metadata['selected_template'] = tmpl
    route_hint = (metadata.get('selected_template_reason') or '').strip()
    msg = f"📝 {TEMPLATE_NAMES.get(tmpl, tmpl)} | {final_model}"
    if route_hint:
        msg += f"\n{route_hint}"
    chunk_chain_text = ' → '.join(LLM_CHUNK_FALLBACK_CHAIN) if LLM_CHUNK_FALLBACK_CHAIN else '(停用額外 fallback)'
    final_chain_text = ' → '.join(LLM_FINAL_FALLBACK_CHAIN) if LLM_FINAL_FALLBACK_CHAIN else '(停用額外 fallback)'
    msg += f"\nChunk fallback: {chunk_chain_text}"
    msg += f"\nFinal fallback: {final_chain_text}"
    if not ALLOW_GOOGLE_DIRECT_FALLBACK:
        msg += "\nGuardrail: Gemini direct fallback 預設停用"
    msg += "\n摘要中..."
    await send_message(msg)

    # Long recording optimization:
    # - DO NOT feed full transcript to the final LLM call
    # - chunk -> cheap model notes -> one final reduce with stronger model
    use_chunk = (tx["duration_sec"] >= 45 * 60) or (len(plain) >= 50000)
    proc['used_chunking'] = bool(use_chunk)
    proc['final_model'] = final_model
    proc['final_fallback_chain'] = LLM_FINAL_FALLBACK_CHAIN[:]
    proc['chunk_fallback_chain'] = LLM_CHUNK_FALLBACK_CHAIN[:]
    proc['llm_policy'] = _describe_llm_policy()

    llm_t0 = time.time()
    try:
        if use_chunk:
            chunk_model = select_chunk_model(tx["duration_sec"]/60, metadata.get("model_pref"))
            await send_message(f"⏳ 長音檔模式：分段整理 → 整合\n  chunks: {chunk_model} → final: {final_model}")

            chunks = _chunk_segments_for_llm(segs, speakers, max_chars=12000, max_minutes=12)
            proc['chunk_count'] = len(chunks)
            proc['chunk_model'] = chunk_model
            await send_message(f"🔪 已切成 {len(chunks)} 段（每段約 10~12 分鐘）")

            # Build template system prompt without embedding transcript
            tmpl_sys, metadata_block, glossary_md = build_system_prompt(tmpl, plain, metadata, speakers)

            import asyncio
            sem = asyncio.Semaphore(3)  # control concurrency to reduce rate-limit risk

            async def _summarize_one(i, ch):
                label = f"{i}/{len(chunks)} {ch['label']}"
                sys_c, usr_c = build_chunk_prompt(ch["text"], label, metadata, speakers)
                async with sem:
                    note = await _call_llm(
                        chunk_model,
                        sys_c,
                        usr_c,
                        send_message,
                        max_tokens=1200,
                        fallback_chain=LLM_CHUNK_FALLBACK_CHAIN,
                        usage_bucket='chunk',
                        proc=proc,
                    )
                return i, label, note

            tasks = [asyncio.create_task(_summarize_one(i, ch)) for i, ch in enumerate(chunks, 1)]

            chunk_notes = [None] * len(chunks)
            done = 0
            chunk_t0 = time.time()
            for fut in asyncio.as_completed(tasks):
                i, label, note = await fut
                chunk_notes[i-1] = f"## {label}\n{note}"
                done += 1
                if done % 4 == 0 or done == len(chunks):
                    await send_message(f"…已完成 {done}/{len(chunks)} 段")
            chunk_elapsed_sec = time.time() - chunk_t0
            proc['chunk_elapsed_sec'] = chunk_elapsed_sec
            proc['chunk_elapsed_text'] = f"{int(chunk_elapsed_sec)} 秒"

            chunk_notes_md = "\n\n".join([x for x in chunk_notes if x])
            reduce_t0 = time.time()
            sys_p, usr_m = build_reduce_prompt(
                tmpl_sys,
                metadata_block,
                chunk_notes_md,
                glossary_md,
                metadata.get('reference_notes', ''),
                metadata=metadata,
            )
            summary = await _call_llm(
                final_model,
                sys_p,
                usr_m,
                send_message,
                max_tokens=6000,
                fallback_chain=LLM_FINAL_FALLBACK_CHAIN,
                usage_bucket='final',
                proc=proc,
            )
            summary_elapsed_sec = time.time() - reduce_t0
            proc['summary_elapsed_sec'] = summary_elapsed_sec
            proc['summary_elapsed_text'] = f"{int(summary_elapsed_sec)} 秒"

        else:
            summary_t0 = time.time()
            sys_p, usr_m = build_summary_prompt(tmpl, plain, metadata, speakers)
            summary = await _call_llm(
                final_model,
                sys_p,
                usr_m,
                send_message,
                max_tokens=6000,
                fallback_chain=LLM_FINAL_FALLBACK_CHAIN,
                usage_bucket='final',
                proc=proc,
            )
            summary_elapsed_sec = time.time() - summary_t0
            proc['summary_elapsed_sec'] = summary_elapsed_sec
            proc['summary_elapsed_text'] = f"{int(summary_elapsed_sec)} 秒"

    except Exception as e:
        await send_message(f"⚠️ LLM 失敗: {e}\n上傳逐字稿...")
        summary = f"## 摘要失敗\n\n{e}"

    # Record total LLM elapsed time
    total_llm_sec = time.time() - llm_t0
    proc['total_elapsed_sec'] = total_llm_sec
    proc['total_elapsed_text'] = f"{int(total_llm_sec)} 秒"

    summary = _strip_emba_transcript_sections(summary, metadata)

    # ===== Step 5: 本地備份 + Obsidian 歸檔 =====
    out = CONFIG["output_dir"]
    out.mkdir(parents=True, exist_ok=True)
    name = _safe_note_name(metadata)
    md_path = out / f"{metadata.get('date','x')}_{name}.md"
    obsidian_md_path = _obsidian_note_path(metadata, _resolve_obsidian_vault_path())
    summary = _append_summary_footer(summary, metadata, audio_path, cache_path if cache_path.exists() else None, md_path, obsidian_md_path)
    local_content = summary
    if _should_include_transcript(metadata) and trans_md.strip():
        local_content += "\n\n---\n\n" + trans_md
    md_path.write_text(local_content, encoding="utf-8")
    obsidian_saved_path = _write_obsidian_markdown(metadata, summary, trans_md)

    # ===== Step 6: Notion =====
    await send_message("📤 上傳 Notion...")
    try:
        if metadata.get("type") == "emba":
            url = upload_emba(metadata, summary, '', tx["duration_sec"])
            label = "課堂摘要庫"
        else:
            metadata["action_items_text"] = _extract_action_items(summary)
            metadata["report_lines"] = _extract_report_lines(summary)
            url = upload_business(metadata, summary, '', tx["duration_sec"])
            label = "商務會談摘要DB"
        msg = f"✅ Done!\n📂 {label}\n🔗 {url}\n💾 本地備份：{md_path}"
        if obsidian_saved_path:
            msg += f"\n💎 Obsidian：{obsidian_saved_path}"
        runtime_lines = _build_model_runtime_lines(proc)
        if runtime_lines:
            msg += "\n" + "\n".join(runtime_lines)
        await send_message(msg)
    except Exception as e:
        msg = f"❌ Notion 失敗: {e}\n💾 本地備份：{md_path}"
        if obsidian_saved_path:
            msg += f"\n💎 Obsidian：{obsidian_saved_path}"
        runtime_lines = _build_model_runtime_lines(proc)
        if runtime_lines:
            msg += "\n" + "\n".join(runtime_lines)
        await send_message(msg)


# ============================================================
# Unified entry point (audio file OR cloud link message)
# ============================================================

async def handle_message(payload, send_message, ask_user, recording_time=None, pending_files=None, **kwargs):
    """Entry point that supports:
    - Telegram audio trigger: payload is local audio_path
    - Telegram text trigger: payload is message text containing Google Drive/OneDrive link

    Notes:
    - For text triggers, we download first, then reuse handle_audio_message.
    - kwargs are accepted for forward-compat with OpenClaw envelopes.
    """
    # Case 1: OpenClaw passes a dict envelope
    if isinstance(payload, dict):
        message_text = payload.get('text') or payload.get('message') or payload.get('message_text')
        audio_path = payload.get('audio_path') or payload.get('file_path')
        caption = payload.get('caption')
        image_paths = _collect_image_paths(payload, **kwargs)
        reference_file_paths = _collect_reference_file_paths(payload, **kwargs)
    else:
        message_text = kwargs.get('message_text') or kwargs.get('text') or ''
        caption = kwargs.get('caption') or ''
        audio_path = payload
        image_paths = _collect_image_paths({}, **kwargs)
        reference_file_paths = _collect_reference_file_paths({}, **kwargs)

    # If it's a real file path, process as audio
    try:
        if isinstance(audio_path, str) and audio_path and Path(audio_path).exists():
            reference_file_paths = [p for p in reference_file_paths if str(p) != str(audio_path)]
            return await handle_audio_message(
                audio_path,
                send_message,
                ask_user,
                recording_time=recording_time,
                pending_files=pending_files,
                message_text=message_text,
                caption=caption,
                image_paths=image_paths,
                reference_file_paths=reference_file_paths,
                reference_urls=_pick_reference_urls((caption or '') + "\n" + (message_text or '')),
            )
    except Exception:
        pass

    # Otherwise treat as message text and look for URLs
    text = ''
    if isinstance(audio_path, str) and audio_path and audio_path.startswith('http'):
        text = audio_path
    else:
        text = (caption or '') + "\n" + (message_text or '')

    audio_url = _pick_audio_url(text)
    reference_urls = _pick_reference_urls(text)
    if not audio_url:
        # Avoid noisy auto-triggers (e.g., discord_message) when the message isn't an audio link.
        # If the user actually wants transcription, they'll paste a link or upload an audio file.
        return

    await send_message("🔗 收到雲端連結，正在下載音訊檔…")
    try:
        dl_path = _download_shared_audio(audio_url, CONFIG["output_dir"] / "downloads")
    except Exception as e:
        await send_message(f"❌ 下載失敗：{e}\n\n請確認連結是『任何知道連結的人都可下載』，或直接把音訊檔丟到 Telegram 給我。")
        return

    await send_message(f"✅ 下載完成：{dl_path.name}\n開始轉錄整理…")
    return await handle_audio_message(
        str(dl_path),
        send_message,
        ask_user,
        recording_time=recording_time,
        pending_files=pending_files,
        message_text=text,
        caption=caption,
        image_paths=image_paths,
        reference_file_paths=reference_file_paths,
        reference_urls=reference_urls,
    )

# ============================================================
# CLI
# ============================================================

if __name__ == "__main__":
    import sys, asyncio
    logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")

    if len(sys.argv) < 2:
        print("單檔: python lecture_pipeline.py <音訊> [YYYY-MM-DD] [HH:MM]")
        print("多檔: python lecture_pipeline.py --merge f1.m4a f2.m4a")
        sys.exit(1)

    async def send(m): print(f"\n💬 {m}")
    async def ask(q): return input(f"\n❓ {q}\n> ")

    if sys.argv[1] == "--merge":
        asyncio.run(handle_audio_message(sys.argv[2], send, ask, pending_files=sys.argv[2:]))
    else:
        rec = None
        if len(sys.argv) >= 3:
            d, t = sys.argv[2], sys.argv[3] if len(sys.argv) >= 4 else "12:00"
            rec = datetime.strptime(f"{d} {t}", "%Y-%m-%d %H:%M")
        asyncio.run(handle_audio_message(sys.argv[1], send, ask, rec))
